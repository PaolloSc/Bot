import os
import re
import json
import time
import logging
import traceback
import unicodedata
from typing import List, Tuple, Optional, Iterable

from seleniumbase import BaseCase
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains

try:
    import pyperclip
except Exception:
    pyperclip = None

try:
    from bs4 import BeautifulSoup
except Exception:
    BeautifulSoup = None

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger("JT_TESTE")


class JTJurisTeste(BaseCase):
    def setUp(self):
        super().setUp()
        self.selectors = self._load_selectors()
        # Estado para sumário e bookmarks por Turma
        self._turma_bookmarks = {}
        self._bookmark_id_counter = 1
        # Estado para Word COM persistente (opcional)
        self._word_app = None
        self._word_persistent = False
        # Flags de ambiente para estabilidade
        self.skip_sumario = os.environ.get("JT_SKIP_SUMARIO", "0") == "1"
        self.disable_clipboard = os.environ.get("JT_DISABLE_CLIPBOARD", "0") == "1"

    def _load_selectors(self, caminho='selectors_jt.json'):
        try:
            with open(caminho, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Erro ao carregar seletores JT: {e}")
            return {"jt": {}}

    # ---------- Sumário e Bookmarks ----------
    def _sanitizar_nome_bookmark(self, nome: str) -> str:
        try:
            base = re.sub(r'[^A-Za-z0-9_]+', '_', nome or '')
            if not base or not base[0].isalpha():
                base = f"BM_{base}"
            return base
        except Exception:
            return "BM_FALLBACK"

    def _inserir_bookmark_no_paragrafo(self, doc: Document, paragrafo, bookmark_name: str):
        try:
            self._bookmark_id_counter += 1
            bid = self._bookmark_id_counter
            p = paragrafo._p
            bm_start = OxmlElement('w:bookmarkStart')
            bm_start.set(qn('w:id'), str(bid))
            bm_start.set(qn('w:name'), bookmark_name)
            p.append(bm_start)
            bm_end = OxmlElement('w:bookmarkEnd')
            bm_end.set(qn('w:id'), str(bid))
            p.append(bm_end)
        except Exception:
            pass

    def _extrair_id_bloco(self, orgao_completo: str) -> Optional[str]:
        try:
            if not orgao_completo:
                return None
            # normaliza removendo acentos/diacríticos
            txt = unicodedata.normalize('NFD', orgao_completo)
            txt = ''.join(c for c in txt if unicodedata.category(c) != 'Mn')
            txt = re.sub(r'[^A-Za-z0-9\sª]+', ' ', txt).lower()

            # Primeiro verificar se tem TRT (ex.: TRT3, TRT 24)
            m_trt = re.search(r'\btrt\s*(\d{1,2})\b', txt)
            if m_trt:
                trt_num = m_trt.group(1)
                # Verificar se também tem turma (ex.: "TRT3 - 1ª Turma")
                m_turma = re.search(r'(\d+)\s*turma', txt)
                if m_turma:
                    turma_num = m_turma.group(1)
                    return f"TRT{trt_num}_{turma_num}ª"  # Ex: TRT3_1ª, TRT24_2ª
                else:
                    return f"TRT{trt_num}"  # Ex: TRT3, TRT24

            # Se não tem TRT mas tem turma, assumir que é só a turma
            m = re.search(r'(\d+)\s*turma', txt)
            if m:
                return f"{m.group(1)}ª"

            # Casos especiais
            if 'conselho superior da justica do trabalho' in txt:
                return 'CSJT'
            if 'tribunal pleno' in txt:
                return 'Pleno'
            if 'orgao especial' in txt:
                return 'Especial'
            return None
        except Exception:
            return None

    def _descricao_por_identificador(self, ident: Optional[str]) -> str:
        try:
            if not ident:
                return 'Processo'
            if ident == 'CSJT':
                return 'Decisão CSJT'
            if ident == 'Pleno':
                return 'Decisão Tribunal Pleno'
            if ident == 'Especial':
                return 'Decisão Órgão Especial'

            # TRT com turma (ex: TRT3_1ª, TRT24_2ª)
            m = re.match(r'TRT(\d+)_(\d+)ª$', ident)
            if m:
                trt_num = m.group(1)
                turma_num = m.group(2)
                return f"TRT {trt_num} - {turma_num}ª Turma"

            # TRT sem turma (ex: TRT3, TRT24)
            if re.match(r'^TRT\d+$', ident):
                num = re.search(r'\d+', ident).group()
                return f"TRT {num} - Acórdãos"

            # Turma simples (ex: 1ª, 2ª)
            if re.match(r'^\d+ª$', ident):
                return f"Acórdão {ident} Turma"

            return 'Processo'
        except Exception:
            return 'Processo'

    def _prepare_document_with_sumario(self, doc_path: str) -> Document:
        try:
            if os.path.exists(doc_path):
                doc = Document(doc_path)
            else:
                # Criar novo doc vazio
                doc = Document()
            # Garantir que exista 'Sumário'
            si = self._buscar_sumario_em_documento(doc)
            if not si:
                if doc.paragraphs:
                    p = doc.paragraphs[0].insert_paragraph_before('Sumário')
                else:
                    p = doc.add_paragraph('Sumário')
                # formatar
                if p.runs:
                    for r in p.runs:
                        r.bold = True
                        r.font.name = 'Arial MT'
                        r.font.size = Pt(9)
                else:
                    run = p.add_run('Sumário')
                    run.bold = True
                    run.font.name = 'Arial MT'
                    run.font.size = Pt(9)
                p.insert_paragraph_after('')
            # Salvar estado
            doc.save(doc_path)
            return doc
        except Exception:
            return Document(doc_path) if os.path.exists(doc_path) else Document()

    def _buscar_sumario_em_documento(self, doc: Document):
        try:
            for i, p in enumerate(doc.paragraphs):
                txt = (p.text or '').strip()
                # normalizar (remover acentos) e comparar de forma robusta
                norm = unicodedata.normalize('NFD', txt)
                norm = ''.join(c for c in norm if unicodedata.category(c) != 'Mn')
                norm = re.sub(r'\s+', ' ', norm).strip().lower()
                if 'sumario' == norm or norm.startswith('sumario'):
                    return {'tipo': 'paragrafo', 'indice': i, 'elemento': p}
            return None
        except Exception:
            return None

    def _remover_paragrafo(self, paragrafo):
        try:
            p = paragrafo._element
            p.getparent().remove(p)
            paragrafo._element = None
        except Exception:
            pass

    def _limpar_sumario_existente(self, doc: Document, sumario_para):
        try:
            # remove linhas seguintes até primeira linha vazia ou até encontrar um heading
            started = False
            removals = []
            for p in list(doc.paragraphs):
                if p is sumario_para:
                    started = True
                    continue
                if not started:
                    continue
                text = (p.text or '').strip()
                # Para quando chegar numa linha vazia ou num heading
                try:
                    is_heading = (p.style and p.style.name and p.style.name.startswith('Heading'))
                except Exception:
                    is_heading = False
                if not text:
                    break
                if is_heading:
                    break
                removals.append(p)
            for rp in removals:
                self._remover_paragrafo(rp)
        except Exception:
            pass

    def _inserir_entrada_sumario(self, doc: Document, sumario_para, label: str, pagina: Optional[int]):
        try:
            novo_para = sumario_para.insert_paragraph_after('')
            run = novo_para.add_run(f"{label}")
            run.font.name = 'Arial MT'
            run.font.size = Pt(9)
            # espaçamento com tabs simples
            novo_para.add_run('\t')
            if pagina is not None:
                pr = novo_para.add_run(str(pagina))
                pr.font.name = 'Arial MT'
                pr.font.size = Pt(9)
            return novo_para
        except Exception:
            return None

    def _inserir_entrada_sumario_pageref(self, doc: Document, sumario_para, label: str, bookmark_name: str):
        try:
            novo_para = sumario_para.insert_paragraph_after('')
            run = novo_para.add_run(f"{label}\t")
            run.font.name = 'Arial MT'
            run.font.size = Pt(9)
            fld = OxmlElement('w:fldSimple')
            fld.set(qn('w:instr'), f"PAGEREF {bookmark_name} \\h")
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = ' '
            r.append(t)
            fld.append(r)
            novo_para._p.append(fld)
            return novo_para
        except Exception:
            return None

    def _get_word_app(self):
        """Obtém instância do Word COM, mantendo-a viva para performance."""
        if getattr(self, '_word_app', None):
            try:
                _ = self._word_app.Version
                return self._word_app
            except Exception:
                self._word_app = None

        try:
            import win32com.client as win32
            try:
                word = win32.GetActiveObject('Word.Application')
            except Exception:
                word = win32.gencache.EnsureDispatch('Word.Application')

            try:
                word.Visible = False
            except Exception:
                pass
            try:
                word.DisplayAlerts = 0
            except Exception:
                pass
            try:
                word.ScreenUpdating = False
            except Exception:
                pass

            self._word_app = word
            self._word_persistent = True
            return word

        except AttributeError:
            try:
                import shutil, tempfile, win32com.client as win32
                shutil.rmtree(os.path.join(tempfile.gettempdir(), 'gen_py'), ignore_errors=True)
                word = win32.Dispatch('Word.Application')
                try:
                    word.Visible = False
                    word.DisplayAlerts = 0
                except Exception:
                    pass
                self._word_app = word
                self._word_persistent = True
                return word
            except Exception:
                return None
        except Exception as e:
            logger.error(f"Erro ao inicializar Word: {e}")
            return None

    def _open_word_app(self) -> bool:
        """Abre (ou reutiliza) uma instância persistente do Word COM.

        Retorna True se a instância estiver disponível.
        """
        try:
            word = self._get_word_app()
            if not word:
                return False
            # marcar como persistente para evitar quit automático
            self._word_persistent = True
            try:
                word.Visible = False
            except Exception:
                pass
            return True
        except Exception:
            return False

    def _close_word_app(self):
        """Fecha a instância persistente do Word COM, se houver."""
        try:
            if getattr(self, '_word_app', None):
                try:
                    self._word_app.Quit()
                except Exception:
                    pass
            self._word_app = None
            self._word_persistent = False
        except Exception:
            pass

    def _ler_paginas_bookmarks_word(self, doc_path: str, bookmark_names: List[str]) -> dict:
        paginas = {}
        try:
            try:
                import win32com.client as win32
            except Exception as e:
                logger.error(f"pywin32 não disponível para leitura de páginas no Word: {e}")
                return {name: None for name in bookmark_names}

            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(doc_path)
            try:
                # Recalcular paginação
                doc.Repaginate()
                # 3 = wdActiveEndPageNumber
                wdActiveEndPageNumber = 3
                for name in bookmark_names:
                    try:
                        if doc.Bookmarks.Exists(name):
                            rng = doc.Bookmarks(name).Range
                            pg = int(rng.Information(wdActiveEndPageNumber))
                            paginas[name] = pg
                        else:
                            paginas[name] = None
                    except Exception:
                        paginas[name] = None
            finally:
                try:
                    doc.Close(False)
                except Exception:
                    pass
                try:
                    word.Quit()
                except Exception:
                    pass
        except Exception as e:
            logger.error(f"Erro ao ler páginas via Word COM: {e}")
            logger.debug(traceback.format_exc())
            paginas = {name: None for name in bookmark_names}
        return paginas
        return paginas

    def _atualizar_sumario_com_paginas(self, doc_path: str, sumario_paginas: dict) -> bool:
        try:
            if not sumario_paginas:
                return False
            doc = Document(doc_path)
            suminfo = self._buscar_sumario_em_documento(doc)
            if not suminfo:
                # criar 'Sumário' se não existir
                doc = self._prepare_document_with_sumario(doc_path)
                suminfo = self._buscar_sumario_em_documento(doc)
                if not suminfo:
                    return False
            sum_para = suminfo['elemento']
            self._limpar_sumario_existente(doc, sum_para)
            ordem_ids = ['CSJT','1ª','2ª','3ª','4ª','5ª','6ª','7ª','8ª','TRT3','TRT24','Pleno','Especial']
            usados = set()
            for ident in ordem_ids:
                if ident in sumario_paginas and sumario_paginas.get(ident) is not None:
                    label = self._descricao_por_identificador(ident)
                    self._inserir_entrada_sumario(doc, sum_para, label, sumario_paginas[ident])
                    usados.add(ident)
            for ident, pg in sumario_paginas.items():
                if ident in usados:
                    continue
                if pg is None:
                    continue
                label = self._descricao_por_identificador(ident)
                self._inserir_entrada_sumario(doc, sum_para, label, pg)
            doc.save(doc_path)
            return True
        except Exception:
            return False

    def _atualizar_sumario_robusto(self, doc_path: str) -> bool:
        """Método robusto para atualizar o sumário baseado em todos os bookmarks do documento usando Word COM."""
        try:
            if not os.path.exists(doc_path):
                logger.warning("Documento não encontrado para atualização do sumário.")
                return False

            # Usar Word COM para ler todos os bookmarks, suas páginas E criar sumário formatado
            try:
                word = self._get_word_app()
                if not word:
                    logger.warning("Não foi possível obter instância do Word COM")
                    return False

                word.Visible = False
                doc_com = word.Documents.Open(os.path.abspath(doc_path))

                # Repaginar documento
                doc_com.Repaginate()

                # Coletar todos os bookmarks e suas páginas
                bookmarks_info = {}
                # Estrutura para acumular páginas por identificador (usada quando não houver bookmarks)
                sumario_data = {}
                wdActiveEndPageNumber = 3

                for bookmark in doc_com.Bookmarks:
                    try:
                        name = bookmark.Name
                        rng = bookmark.Range
                        page = int(rng.Information(wdActiveEndPageNumber))
                        bookmarks_info[name] = page
                    except Exception:
                        continue

                logger.info(f"Encontrados {len(bookmarks_info)} bookmarks no documento")

            except Exception as e:
                logger.error(f"Erro ao ler bookmarks via Word COM: {e}")
                if 'doc_com' in locals():
                    try:
                        doc_com.Close(False)
                        if not getattr(self, '_word_persistent', False):
                            word.Quit()
                    except:
                        pass
                return False
            
            if not bookmarks_info:
                logger.warning("Nenhum bookmark encontrado no documento")
                return False
            
            # Organizar dados do sumário
            sumario_data = {}  # {turma/orgao: menor_pagina}
            
            # Processar bookmarks BM_TURMA_* (TRTs do JT)
            for bm_name, page in bookmarks_info.items():
                if bm_name.startswith('BM_TURMA_'):
                    ident = bm_name.replace('BM_TURMA_', '')
                    if ident not in sumario_data or page < sumario_data[ident]:
                        sumario_data[ident] = page
            
            # Processar bookmarks BM_PROC_* (processos TST)
            # Separar CSJT (5.90.0000) das demais turmas
            all_proc_bookmarks = [(bm, page) for bm, page in bookmarks_info.items() if bm.startswith('BM_PROC_')]
            
            # Identificar processos CSJT
            csjt_bookmarks = []
            proc_bookmarks = []
            
            for bm, page in all_proc_bookmarks:
                if '5_90_0000' in bm:  # Processos CSJT têm 5.90.0000
                    csjt_bookmarks.append((bm, page))
                else:
                    proc_bookmarks.append((bm, page))
            
            # Processar CSJT
            if csjt_bookmarks:
                # Pegar a primeira página dos processos CSJT
                csjt_pages = [page for bm, page in csjt_bookmarks]
                if csjt_pages:
                    sumario_data['CSJT'] = min(csjt_pages)
            
            if proc_bookmarks:
                # Ordenar por página
                proc_bookmarks.sort(key=lambda x: x[1])
                
                # Estimar turmas baseado na distribuição de páginas
                turmas_tst = ['1ª', '2ª', '3ª', '4ª', '5ª', '6ª', '7ª', '8ª']
                
                # Dividir processos em grupos e atribuir às turmas
                if len(proc_bookmarks) >= len(turmas_tst):
                    processos_por_turma = len(proc_bookmarks) // len(turmas_tst)
                    for i, turma in enumerate(turmas_tst):
                        inicio = i * processos_por_turma
                        if inicio < len(proc_bookmarks):
                            page = proc_bookmarks[inicio][1]
                            if turma not in sumario_data or page < sumario_data[turma]:
                                sumario_data[turma] = page
                else:
                    for i, (bm, page) in enumerate(proc_bookmarks):
                        if i < len(turmas_tst):
                            turma = turmas_tst[i]
                            if turma not in sumario_data or page < sumario_data[turma]:
                                sumario_data[turma] = page
            
            if not sumario_data:
                logger.warning("Não há dados para atualizar o sumário")
                try:
                    doc_com.Close(False)
                    if not getattr(self, '_word_persistent', False):
                        word.Quit()
                except:
                    pass
                return False

            # Agora vamos criar o sumário usando a API COM do Word para formatação adequada
            try:
                # Procurar parágrafo "Sumário" no documento
                sumario_range = None
                for para in doc_com.Paragraphs:
                    texto_para = para.Range.Text.strip().lower()
                    # Normalizar removendo acentos
                    import unicodedata
                    texto_norm = unicodedata.normalize('NFD', texto_para)
                    texto_norm = ''.join([c for c in texto_norm if not unicodedata.combining(c)])

                    if 'sumario' in texto_norm or texto_norm.startswith('sum'):
                        sumario_range = para.Range
                        # Limpar conteúdo após "Sumário" (entradas antigas)
                        # Deletar próximos parágrafos que parecem ser entradas de sumário
                        idx_atual = para.Range.Paragraphs(1).Index
                        paras_deletar = []

                        for i in range(idx_atual + 1, min(idx_atual + 25, doc_com.Paragraphs.Count + 1)):
                            try:
                                p = doc_com.Paragraphs(i)
                                txt = p.Range.Text.strip()
                                # Se é entrada de sumário (tem "Acórdão", "Decisão", "TRT" ou pontos seguidos de número)
                                if (txt and ('Acórdão' in txt or 'Decisão' in txt or 'TRT' in txt or
                                    re.search(r'\.*\s*\d+\s*$', txt) or '...' in txt)):
                                    paras_deletar.append(i)
                                elif txt and len(txt) > 100:  # Conteúdo normal do documento
                                    break
                                elif not txt:  # Linha vazia - pode fazer parte do sumário
                                    paras_deletar.append(i)
                            except:
                                break

                        # Deletar de trás para frente para não afetar índices
                        for idx in reversed(paras_deletar):
                            try:
                                doc_com.Paragraphs(idx).Range.Delete()
                            except:
                                pass
                        break

                # Se não encontrou, criar no início do documento
                if sumario_range is None:
                    rng = doc_com.Range(0, 0)
                    rng.InsertAfter("Sumário\n")
                    rng.Font.Bold = True
                    rng.Font.Size = 14
                    rng.Font.Name = "Arial"
                    sumario_range = doc_com.Paragraphs(1).Range

                # Posicionar cursor após o título "Sumário"
                insert_range = sumario_range.Paragraphs(1).Range
                insert_range.Collapse(0)  # wdCollapseEnd = 0
                insert_range.Move(5, 1)  # wdParagraph = 5, move 1 paragraph down

                # Preparar dados do sumário em ordem
                ordem = ['CSJT', 'Pleno', 'Especial', '1ª', '2ª', '3ª', '4ª', '5ª', '6ª', '7ª', '8ª']
                # Adicionar TRTs em ordem numérica
                trts_list = sorted([k for k in sumario_data.keys() if k.startswith('TRT')],
                             key=lambda x: int(re.search(r'\d+', x).group()) if re.search(r'\d+', x) else 999)
                ordem.extend(trts_list)
                ordem.extend([k for k in sumario_data.keys() if k not in ordem])

                # Mapear bookmarks por identificador
                bookmark_map = {}  # {ident: bookmark_name}
                for bm_name in bookmarks_info.keys():
                    if bm_name.startswith('BM_TURMA_'):
                        ident = bm_name.replace('BM_TURMA_', '')
                        bookmark_map[ident] = bm_name

                # Inserir entradas do sumário com hyperlinks
                entradas_inseridas = 0
                for ident in ordem:
                    if ident not in sumario_data:
                        continue

                    # Mapear descrição
                    if ident == 'CSJT':
                        descricao = 'Decisão CSJT'
                    elif ident == 'Pleno':
                        descricao = 'Decisão Tribunal Pleno'
                    elif ident == 'Especial':
                        descricao = 'Decisão Órgão Especial'
                    elif re.match(r'^\d+ª$', ident):
                        descricao = f"Acórdão {ident} Turma"
                    elif ident.startswith('TRT'):
                        num = re.search(r'\d+', ident).group()
                        if len(num) == 1:
                            descricao = f"TRT{num} Acórdão"
                        else:
                            descricao = f"TRT {num} Acórdão"
                    else:
                        descricao = 'Processo'

                    page = sumario_data[ident]
                    bookmark_name = bookmark_map.get(ident, '')

                    # Inserir nova linha
                    insert_range.InsertAfter(descricao)

                    # Criar hyperlink se temos bookmark
                    if bookmark_name and doc_com.Bookmarks.Exists(bookmark_name):
                        # Selecionar o texto que acabamos de inserir
                        start_pos = insert_range.End - len(descricao)
                        link_range = doc_com.Range(start_pos, insert_range.End)

                        # Criar hyperlink para o bookmark
                        try:
                            doc_com.Hyperlinks.Add(
                                Anchor=link_range,
                                SubAddress=bookmark_name,
                                ScreenTip=f"Ir para {descricao}"
                            )
                        except:
                            pass  # Se falhar, continua sem hyperlink

                    # Adicionar tabulação e número de página
                    insert_range.InsertAfter("\t" + str(page) + "\n")

                    # Formatar o parágrafo
                    para_atual = insert_range.Paragraphs(insert_range.Paragraphs.Count)
                    para_atual.Range.Font.Name = "Arial"
                    para_atual.Range.Font.Size = 11
                    para_atual.Range.Font.Bold = False

                    # Adicionar tabulação com preenchimento de pontos
                    try:
                        para_atual.TabStops.ClearAll()
                        # Adicionar tab stop na posição 450 (aproximadamente margem direita)
                        # wdAlignTabRight = 2, wdTabLeaderDots = 1
                        para_atual.TabStops.Add(Position=450, Alignment=2, Leader=1)
                    except:
                        pass

                    entradas_inseridas += 1

                # Salvar e fechar
                doc_com.Save()
                doc_com.Close(False)
                if not getattr(self, '_word_persistent', False):
                    word.Quit()

                logger.info(f"✅ Sumário atualizado com {entradas_inseridas} entradas formatadas e com hyperlinks")
                return True

            except Exception as e:
                logger.error(f"Erro ao criar sumário formatado: {e}")
                logger.error(traceback.format_exc())
                try:
                    doc_com.Close(False)
                    if not getattr(self, '_word_persistent', False):
                        word.Quit()
                except:
                    pass
                return False
            
        except Exception as e:
            logger.error(f"Erro ao atualizar sumário robusto: {e}")
            logger.error(traceback.format_exc())
            return False
        # Fechar instância persistente do Word COM (se aberta)
        try:
            if getattr(self, '_word_persistent', False):
                self._close_word_app()
                logger.info("Word COM persistente fechado.")
        except Exception:
            pass
    
    def _atualizar_sumario_jt_simples(self, doc_path: str) -> bool:
        """
        Atualiza o sumário do documento JT de forma simples (baseado no bot TST).
        Usa bookmarks para encontrar as páginas dos TRTs.
        """
        try:
            if not os.path.exists(doc_path):
                logger.warning("Documento não encontrado para atualização do sumário.")
                return False

            # Coletar bookmarks e suas páginas usando win32com
            if not hasattr(self, '_turma_bookmarks') or not self._turma_bookmarks:
                logger.warning("Nenhum bookmark de turma encontrado em _turma_bookmarks")
                return False

            logger.info(f"Coletando páginas dos bookmarks: {list(self._turma_bookmarks.keys())}")

            # Ler páginas usando win32com
            ident_to_bm = dict(self._turma_bookmarks)
            paginas = self._ler_paginas_bookmarks_word(doc_path, list(ident_to_bm.values()))

            sumario_paginas = {}
            for ident, bm in ident_to_bm.items():
                pg = paginas.get(bm)
                if pg:
                    sumario_paginas[ident] = pg
                    logger.info(f"  {ident} → página {pg}")

            if not sumario_paginas:
                logger.warning("Nenhuma página encontrada para os bookmarks")
                return False

            # Atualizar sumário usando python-docx
            doc = Document(doc_path)
            suminfo = self._buscar_sumario_em_documento(doc)

            if not suminfo:
                logger.warning("Parágrafo 'Sumário' não encontrado")
                return False

            sum_para = suminfo['elemento']

            # Limpar entradas antigas
            self._limpar_sumario_existente(doc, sum_para)

            # Separar identificadores por categoria
            trt3_turmas = []
            trt24_turmas = []
            outros = []

            for ident in sumario_paginas.keys():
                # TRT3 com turma (ex: TRT3_1ª, TRT3_2ª)
                if ident.startswith('TRT3_'):
                    trt3_turmas.append(ident)
                # TRT24 com turma (ex: TRT24_1ª, TRT24_2ª)
                elif ident.startswith('TRT24_'):
                    trt24_turmas.append(ident)
                # TRT3 ou TRT24 genérico
                elif ident in ['TRT3', 'TRT24']:
                    outros.append(ident)
                # Turmas sem TRT especificado (1ª, 2ª, etc.)
                elif re.match(r'^\d+ª$', ident):
                    trt3_turmas.append(ident)  # Assumir TRT3 por padrão
                else:
                    outros.append(ident)

            # Ordenar turmas numericamente
            def ordenar_turma(t):
                # Para TRT3_1ª ou TRT24_1ª
                m = re.match(r'TRT\d+_(\d+)ª', t)
                if m:
                    return int(m.group(1))
                # Para 1ª, 2ª
                m = re.match(r'(\d+)ª', t)
                return int(m.group(1)) if m else 999

            trt3_turmas.sort(key=ordenar_turma)
            trt24_turmas.sort(key=ordenar_turma)

            # Inserir no sumário em ordem
            # 1. TRT3 e suas turmas
            logger.info(f"📝 Separando entradas do sumário:")
            logger.info(f"  - TRT3: {trt3_turmas}")
            logger.info(f"  - TRT24: {trt24_turmas}")
            logger.info(f"  - Outros: {outros}")

            logger.info(f"Inserindo TRT3: {len(trt3_turmas)} turmas")
            for turma in trt3_turmas:
                label = self._descricao_por_identificador(turma)
                entrada = f"{label}\t{sumario_paginas[turma]}"
                logger.info(f"  ✓ Inserindo: {entrada}")
                self._inserir_entrada_sumario_simples(doc, sum_para, entrada)

            # 2. TRT24 e suas turmas
            logger.info(f"Inserindo TRT24: {len(trt24_turmas)} turmas")
            for turma in trt24_turmas:
                label = self._descricao_por_identificador(turma)
                entrada = f"{label}\t{sumario_paginas[turma]}"
                logger.info(f"  ✓ Inserindo: {entrada}")
                self._inserir_entrada_sumario_simples(doc, sum_para, entrada)

            # 3. Outros identificadores (TRT3/TRT24 genéricos, CSJT, etc.)
            logger.info(f"Inserindo outros: {len(outros)} identificadores")
            for ident in outros:
                label = self._descricao_por_identificador(ident)
                entrada = f"{label}\t{sumario_paginas[ident]}"
                logger.info(f"  ✓ Inserindo: {entrada}")
                self._inserir_entrada_sumario_simples(doc, sum_para, entrada)

            doc.save(doc_path)
            logger.info(f"✅ Sumário atualizado com {len(sumario_paginas)} entradas (turmas separadas por TRT)")
            return True

        except Exception as e:
            logger.error(f"Erro ao atualizar sumário simples: {e}")
            logger.error(traceback.format_exc())
            return False

    def _inserir_entrada_sumario_simples(self, doc: Document, sumario_para, entrada: str):
        """Insere uma entrada no sumário com formatação simples."""
        try:
            novo_para = sumario_para.insert_paragraph_after('')
            run = novo_para.add_run(entrada)
            run.font.name = 'Arial MT'
            run.font.size = Pt(9)

            # Configurar tab stop com pontos
            try:
                from docx.shared import Inches
                novo_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), alignment=2, leader=1)  # 2=right, 1=dots
            except:
                pass

            return novo_para
        except Exception as e:
            logger.error(f"Erro ao inserir entrada: {e}")
            return None

    def _atualizar_sumario_jt_word_com(self, doc_path: str) -> bool:
        """
        Atualiza o sumário do documento JT usando Word COM.
        Insere entradas dos TRTs extraídos logo após o título 'SUMÁRIO',
        preservando o conteúdo existente (Secretaria, Pauta, etc.).
        """
        try:
            if not os.path.exists(doc_path):
                logger.warning("Documento não encontrado para atualização do sumário.")
                return False

            # Caminho definitivo: construir entradas via PAGEREF e forçar atualização via helper externo
            try:
                if getattr(self, '_turma_bookmarks', None):
                    ident_to_bm = dict(self._turma_bookmarks)
                    # 1) Insere entradas com campos PAGEREF (python-docx) — sem depender de COM agora
                    ok_pageref = self._atualizar_sumario_com_pageref(doc_path, ident_to_bm)
                    if ok_pageref:
                        # 2) Tenta atualizar campos/TOC usando processo externo robusto (pywin32)
                        try:
                            self._forcar_atualizacao_campos_via_word(doc_path)
                        except Exception as e:
                            logger.warning(f"Atualizador externo falhou (prosseguindo com doc salvo): {e}")
                        return True
            except Exception as e:
                logger.debug(f"Rota PAGEREF+helper falhou, mantendo fluxo antigo: {e}")

            # Abrir Word usando função auxiliar
            word = self._get_word_app()
            if not word:
                logger.warning("Não foi possível obter instância do Word COM")
                return False

            word.Visible = False
            doc_com = word.Documents.Open(os.path.abspath(doc_path))

            try:
                # Repaginar documento
                doc_com.Repaginate()

                # Coletar todos os bookmarks e suas páginas
                bookmarks_info = {}
                sumario_data = {}
                wdActiveEndPageNumber = 3

                for bookmark in doc_com.Bookmarks:
                    try:
                        name = bookmark.Name
                        rng = bookmark.Range
                        page = int(rng.Information(wdActiveEndPageNumber))
                        bookmarks_info[name] = page
                    except Exception:
                        try:
                            if not os.path.exists(doc_path):
                                logger.warning("Documento não encontrado para atualizar sumário.")
                                return False

                            # Tentar usar win32com abrindo/fechando o Word por chamada (mais robusto que instância persistente)
                            try:
                                import win32com.client as win32
                                import pywintypes
                            except Exception as e:
                                logger.warning(f"pywin32 não disponível para atualização JT via COM: {e}")
                                # Fallback para método simples baseado em python-docx
                                return self._atualizar_sumario_jt_simples(doc_path)

                            abs_path = os.path.abspath(doc_path)
                            last_exc = None
                            for attempt in range(1, 4):
                                try:
                                    word = win32.gencache.EnsureDispatch('Word.Application')
                                    word.Visible = False
                                    doc_com = word.Documents.Open(abs_path)
                                    try:
                                        # Repaginar e atualizar TOC/fields
                                        try:
                                            doc_com.Repaginate()
                                        except Exception:
                                            pass

                                        try:
                                            toc_count = doc_com.TablesOfContents.Count
                                        except Exception:
                                            toc_count = 0

                                        if toc_count and toc_count > 0:
                                            for i in range(1, toc_count + 1):
                                                try:
                                                    doc_com.TablesOfContents(i).Update()
                                                except Exception:
                                                    pass

                                        try:
                                            doc_com.Fields.Update()
                                        except Exception:
                                            pass

                                        doc_com.Save()
                                    finally:
                                        try:
                                            doc_com.Close(False)
                                        except Exception:
                                            pass
                                        try:
                                            word.Quit()
                                        except Exception:
                                            pass

                                    logger.info("✅ Sumário JT atualizado via Word COM.")
                                    return True
                                except pywintypes.com_error as ce:
                                    last_exc = ce
                                    # Erros comuns: RPC server not available (-2147023174) or call rejected (-2147418111)
                                    code = int(ce.args[0]) if ce.args and isinstance(ce.args[0], int) else None
                                    logger.warning(f"Tentativa {attempt}: erro COM ao atualizar sumário JT: {ce}")
                                    # pequena espera exponencial antes de retry
                                    time.sleep(0.8 * (2 ** (attempt - 1)))
                                    continue
                                except Exception as e:
                                    last_exc = e
                                    logger.debug(f"Tentativa {attempt}: exceção inesperada ao usar COM: {e}")
                                    time.sleep(0.5)
                                    continue

                            logger.error(f"Falha ao atualizar sumário via Word COM após retries: {last_exc}")
                            # Fallback para método simples baseado em python-docx
                            try:
                                return self._atualizar_sumario_jt_simples(doc_path)
                            except Exception:
                                logger.error("Fallback simples também falhou ao atualizar sumário.")
                                return False
                        except Exception as e:
                            logger.error(f"Erro geral ao atualizar sumário: {e}")
                            logger.error(traceback.format_exc())
                            return False
                    logger.info("Nenhum bookmark BM_TURMA_ encontrado, procurando por padrões de TRT nos parágrafos...")

                    # Procurar por parágrafos que contenham TRT3 ou TRT24
                    for i in range(1, doc_com.Paragraphs.Count + 1):
                        try:
                            para = doc_com.Paragraphs(i)
                            texto = para.Range.Text.strip()

                            # Procurar por TRT3 ou TRT24
                            if re.search(r'TRT\s*3\b', texto, re.IGNORECASE):
                                page = int(para.Range.Information(wdActiveEndPageNumber))
                                if 'TRT3' not in sumario_data or page < sumario_data['TRT3']:
                                    sumario_data['TRT3'] = page
                                    logger.info(f"✓ Encontrado TRT3 no parágrafo {i}, página {page}")

                            if re.search(r'TRT\s*24\b', texto, re.IGNORECASE):
                                page = int(para.Range.Information(wdActiveEndPageNumber))
                                if 'TRT24' not in sumario_data or page < sumario_data['TRT24']:
                                    sumario_data['TRT24'] = page
                                    logger.info(f"✓ Encontrado TRT24 no parágrafo {i}, página {page}")
                        except:
                            continue

                if not sumario_data:
                    logger.warning("Nenhum dado de TRT encontrado no documento (nem bookmarks nem parágrafos)")
                    doc_com.Close(False)
                    if not getattr(self, '_word_persistent', False):
                        word.Quit()
                    return False

                # Procurar o parágrafo "SUMÁRIO" (em maiúsculas)
                sumario_para_idx = None
                import unicodedata

                for i in range(1, doc_com.Paragraphs.Count + 1):
                    try:
                        para = doc_com.Paragraphs(i)
                        texto = para.Range.Text.strip().upper()
                        # Normalizar removendo acentos
                        texto_norm = unicodedata.normalize('NFD', texto)
                        texto_norm = ''.join([c for c in texto_norm if not unicodedata.combining(c)])

                        if texto_norm == 'SUMARIO' or texto == 'SUMÁRIO':
                            sumario_para_idx = i
                            logger.info(f"✓ SUMÁRIO encontrado no parágrafo {i}")
                            break
                    except:
                        continue

                if sumario_para_idx is None:
                    logger.warning("Parágrafo 'SUMÁRIO' não encontrado no documento")
                    doc_com.Close(False)
                    if not getattr(self, '_word_persistent', False):
                        word.Quit()
                    return False

                # Remover entradas antigas dos TRTs (se existirem)
                # Procurar e deletar linhas que contenham "TRT3" ou "TRT24" após o SUMÁRIO
                paras_deletar = []
                for i in range(sumario_para_idx + 1, min(sumario_para_idx + 50, doc_com.Paragraphs.Count + 1)):
                    try:
                        para = doc_com.Paragraphs(i)
                        texto = para.Range.Text.strip()
                        # Se contém TRT3, TRT24, ou padrões de turmas do TRT
                        if texto and ('TRT3' in texto or 'TRT24' in texto or
                                     re.search(r'TRT\s*\d+.*Turma', texto)):
                            paras_deletar.append(i)
                    except:
                        break

                # Deletar de trás para frente
                for idx in reversed(paras_deletar):
                    try:
                        doc_com.Paragraphs(idx).Range.Delete()
                        logger.info(f"Removida entrada antiga: parágrafo {idx}")
                    except:
                        pass

                # Preparar ordem de inserção
                ordem = ['TRT3', 'TRT24']

                # Mapear bookmarks
                bookmark_map = {}
                for bm_name in bookmarks_info.keys():
                    if bm_name.startswith('BM_TURMA_'):
                        ident = bm_name.replace('BM_TURMA_', '')
                        bookmark_map[ident] = bm_name

                # Posicionar após o título SUMÁRIO para inserir as entradas
                # Inserir logo após o parágrafo SUMÁRIO
                insert_pos = doc_com.Paragraphs(sumario_para_idx).Range.End

                entradas_inseridas = 0
                for ident in ordem:
                    if ident not in sumario_data:
                        continue

                    page = sumario_data[ident]
                    bookmark_name = bookmark_map.get(ident, '')

                    # Criar descrição
                    if ident.startswith('TRT'):
                        num = re.search(r'\d+', ident).group()
                        descricao = f"TRT {num} - Acórdãos"
                    else:
                        descricao = f"{ident}"

                    # Inserir novo parágrafo
                    rng = doc_com.Range(insert_pos, insert_pos)
                    rng.InsertAfter(f"\n{descricao}\t{page}")

                    # Formatar o parágrafo recém-inserido
                    # Encontrar o parágrafo que acabamos de criar
                    new_para_idx = sumario_para_idx + entradas_inseridas + 1
                    if new_para_idx <= doc_com.Paragraphs.Count:
                        new_para = doc_com.Paragraphs(new_para_idx)
                        new_para.Range.Font.Name = "Arial"
                        new_para.Range.Font.Size = 11
                        new_para.Range.Font.Bold = False

                        # Configurar tab stops com pontos
                        try:
                            new_para.TabStops.ClearAll()
                            # wdAlignTabRight = 2, wdTabLeaderDots = 1
                            new_para.TabStops.Add(Position=450, Alignment=2, Leader=1)
                        except:
                            pass

                        # Criar hyperlink se temos bookmark
                        if bookmark_name and doc_com.Bookmarks.Exists(bookmark_name):
                            try:
                                # Selecionar apenas o texto da descrição (antes do tab)
                                texto_completo = new_para.Range.Text
                                pos_tab = texto_completo.find('\t')
                                if pos_tab > 0:
                                    link_start = new_para.Range.Start
                                    link_end = link_start + pos_tab
                                    link_range = doc_com.Range(link_start, link_end)

                                    doc_com.Hyperlinks.Add(
                                        Anchor=link_range,
                                        SubAddress=bookmark_name,
                                        ScreenTip=f"Ir para {descricao}"
                                    )
                            except Exception as e:
                                logger.debug(f"Não foi possível criar hyperlink: {e}")

                    entradas_inseridas += 1
                    # Atualizar posição para próxima inserção
                    insert_pos = doc_com.Paragraphs(sumario_para_idx + entradas_inseridas).Range.End

                # Salvar e fechar
                doc_com.Save()
                doc_com.Close(False)
                if not getattr(self, '_word_persistent', False):
                    word.Quit()

                logger.info(f"✅ Sumário atualizado: {entradas_inseridas} entradas de TRT inseridas")
                return True

            except Exception as e:
                logger.error(f"Erro ao atualizar sumário JT: {e}")
                logger.error(traceback.format_exc())
                try:
                    doc_com.Close(False)
                    if not getattr(self, '_word_persistent', False):
                        word.Quit()
                except:
                    pass
                return False

        except Exception as e:
            logger.error(f"Erro geral ao atualizar sumário: {e}")
            logger.error(traceback.format_exc())
            return False

    def _forcar_atualizacao_campos_via_word(self, doc_path: str) -> bool:
        """Executa o atualizador externo (COM) em subprocesso para repaginar/atualizar TOC e campos.
        Retorna True em caso de sucesso; se falhar, loga e retorna False sem interromper o fluxo."""
        try:
            import subprocess, sys
            updater = os.path.join(os.path.dirname(__file__), 'word_toc_updater.py')
            if not os.path.exists(updater):
                logger.warning("Helper 'word_toc_updater.py' não encontrado; pulando atualização COM externa.")
                return False
            cmd = [sys.executable, updater, os.path.abspath(doc_path)]
            res = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
            if res.returncode == 0:
                logger.info("Campos/TOC atualizados via helper externo (COM).")
                return True
            logger.warning(f"Helper externo retornou código {res.returncode}. stderr: {res.stderr.strip()}")
            return False
        except Exception as e:
            logger.warning(f"Falha ao executar helper externo: {e}")
            return False

    def _atualizar_sumario_com_pageref(self, doc_path: str, ident_to_bm: dict) -> bool:
        try:
            doc = Document(doc_path)
            suminfo = self._buscar_sumario_em_documento(doc)
            if not suminfo:
                # criar 'Sumário' se não existir
                doc = self._prepare_document_with_sumario(doc_path)
                suminfo = self._buscar_sumario_em_documento(doc)
                if not suminfo:
                    return False
            sum_para = suminfo['elemento']
            self._limpar_sumario_existente(doc, sum_para)
            ordem_ids = ['CSJT','1ª','2ª','3ª','4ª','5ª','6ª','7ª','8ª','TRT3','TRT24','Pleno','Especial']
            usados = set()
            for ident in ordem_ids:
                bm = ident_to_bm.get(ident)
                if not bm:
                    continue
                label = self._descricao_por_identificador(ident)
                self._inserir_entrada_sumario_pageref(doc, sum_para, label, bm)
                usados.add(ident)
            for ident, bm in ident_to_bm.items():
                if ident in usados:
                    continue
                label = self._descricao_por_identificador(ident)
                self._inserir_entrada_sumario_pageref(doc, sum_para, label, bm)
            doc.save(doc_path)
            return True
        except Exception:
            return False

    # ---------- Utilitários ----------
    def _scroll_center(self, el):
        try:
            self.driver.execute_script("arguments[0].scrollIntoView({block:'center'});", el)
        except Exception:
            pass

    def _hover(self, el):
        try:
            ActionChains(self.driver).move_to_element(el).perform()
        except Exception:
            pass

    def _scroll_by(self, fraction=0.85):
        try:
            self.driver.execute_script("window.scrollBy(0, Math.floor(window.innerHeight*arguments[0]));", fraction)
        except Exception:
            try:
                self.driver.execute_script("window.scrollBy(0, 600);")
            except Exception:
                pass

    def _scroll_to_bottom(self, tries: int = 3):
        for _ in range(max(1, tries)):
            try:
                self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            except Exception:
                pass
            time.sleep(0.4)

    def _find_scrollable_ancestor(self, element):
        try:
            js = """
            function findScrollable(el){
              while(el && el !== document.body){
                const cs = window.getComputedStyle(el);
                const oy = (cs.overflowY || cs.overflow || '').toLowerCase();
                if ((el.scrollHeight - el.clientHeight) > 8 && (/auto|scroll/).test(oy)) return el;
                el = el.parentElement;
              }
              return document.scrollingElement || document.documentElement || document.body;
            }
            return findScrollable(arguments[0]);
            """
            return self.driver.execute_script(js, element)
        except Exception:
            return None

    def _scroll_node_to_bottom(self, node, tries: int = 3):
        for _ in range(max(1, tries)):
            try:
                if node is not None and node.tag_name.lower() not in ('html','body'):
                    self.driver.execute_script("arguments[0].scrollTop = arguments[0].scrollHeight;", node)
                else:
                    self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            except Exception:
                pass
            time.sleep(0.35)

    def _reveal_all_cards_on_page(self, max_scrolls: int = 20, settle_cycles: int = 2) -> int:
        prev = -1
        stable = 0
        scrolls = 0
        while scrolls < max_scrolls and stable < settle_cycles:
            cards = self._buscar_cartoes()
            count = len(cards)
            if count <= prev:
                stable += 1
            else:
                stable = 0
            prev = count
            # tentar rolar no ancestral rolável do último card; se não houver, rolar a janela
            try:
                target = cards[-1] if cards else None
            except Exception:
                target = None
            container = None
            if target is not None:
                self._scroll_center(target)
                container = self._find_scrollable_ancestor(target)
            self._scroll_node_to_bottom(container, tries=1)
            scrolls += 1
        return prev if prev >= 0 else 0

    def _ajustar_itens_por_pagina(self, valor: str = '10') -> bool:
        try:
            logger.info(f"📊 Configurando paginação para {valor} itens por página...")
            sel_pag = self.selectors.get('jt', {}).get('resultados', {}).get('paginacao', {})
            labels = sel_pag.get('rows_dropdown_label', [])
            opts10 = sel_pag.get('rows_option_10', [])

            # Rolar ao fundo para garantir visibilidade do paginador
            logger.info("  📜 Rolando até o final da página...")
            self._scroll_to_bottom(tries=3)
            time.sleep(0.8)

            alvo = None
            xpath_encontrado = None
            logger.info(f"  🔍 Procurando dropdown em {len(labels)} seletores...")
            for idx, xp in enumerate(labels, 1):
                try:
                    spans = self.driver.find_elements(By.XPATH, xp)
                    logger.debug(f"    Seletor {idx}: {xp} -> {len(spans)} elementos")
                    for s in spans:
                        try:
                            if s.is_displayed():
                                alvo = s
                                xpath_encontrado = xp
                                logger.info(f"  ✓ Dropdown encontrado com seletor {idx}: {xp[:80]}...")
                                break
                        except Exception:
                            continue
                    if alvo:
                        break
                except Exception as e:
                    logger.debug(f"    Erro no seletor {idx}: {e}")
                    continue

            if not alvo:
                logger.warning("⚠ Dropdown de itens por página não localizado após testar todos os seletores.")
                logger.warning(f"   Tentou {len(labels)} seletores diferentes")
                return False

            # Abrir dropdown
            logger.info(f"  🖱️ Clicando no dropdown...")
            self._scroll_center(alvo)
            self._hover(alvo)

            # Tentar clicar de várias formas
            clicked = False
            try:
                alvo.click()
                clicked = True
            except Exception:
                if self._safe_js_click(alvo):
                    clicked = True
                else:
                    try:
                        self.driver.execute_script("arguments[0].click();", alvo)
                        clicked = True
                    except Exception:
                        pass

            if not clicked:
                logger.warning("  ⚠ Não foi possível clicar no dropdown")
                return False

            time.sleep(0.8)

            # Verificar se dropdown CORRETO abriu (deve conter números, não texto)
            listbox_aberto = False
            listbox_correto = False
            try:
                listboxes = self.driver.find_elements(By.XPATH, "//ul[@role='listbox']")
                for lb in listboxes:
                    if lb.is_displayed():
                        listbox_aberto = True
                        # Verifica se contém opções numéricas (5, 10, etc) em vez de texto
                        try:
                            opcoes_texto = [li.text.strip() for li in lb.find_elements(By.TAG_NAME, 'li') if li.is_displayed()]
                            # Verifica se pelo menos uma opção é numérica
                            tem_numeros = any(opt.isdigit() for opt in opcoes_texto if opt)
                            if tem_numeros:
                                listbox_correto = True
                                logger.info(f"  ✓ Dropdown CORRETO abriu: {opcoes_texto}")
                                break
                            else:
                                logger.warning(f"  ⚠ Dropdown ERRADO abriu (ordenação): {opcoes_texto}")
                        except Exception:
                            pass
            except Exception:
                pass

            if not listbox_aberto:
                logger.warning("  ⚠ Nenhum dropdown abriu. Tentando clicar novamente...")
                time.sleep(0.3)
                try:
                    self.driver.execute_script("arguments[0].click();", alvo)
                    time.sleep(0.6)
                except Exception:
                    pass
            elif listbox_aberto and not listbox_correto:
                logger.warning("  ⚠ Dropdown ERRADO aberto! Fechando e procurando dropdown correto...")
                # Fecha o dropdown errado clicando em ESC
                try:
                    self.driver.find_element(By.TAG_NAME, 'body').send_keys(Keys.ESCAPE)
                    time.sleep(0.4)
                except Exception:
                    pass

                # Procura especificamente pelo dropdown dentro do paginator
                logger.info("  🔍 Procurando dropdown dentro de p-paginator...")
                try:
                    paginator_dropdowns = self.driver.find_elements(
                        By.XPATH,
                        "//div[contains(@class, 'p-paginator')]//span[contains(@class, 'p-dropdown-label')]"
                    )
                    if paginator_dropdowns:
                        alvo = paginator_dropdowns[0]
                        logger.info(f"  ✓ Encontrado dropdown em p-paginator: {alvo.text}")
                        self._scroll_center(alvo)
                        self.driver.execute_script("arguments[0].click();", alvo)
                        time.sleep(0.6)
                except Exception as e:
                    logger.warning(f"  ⚠ Erro ao procurar em p-paginator: {e}")
                    return False

            # Selecionar 10
            logger.info(f"  🔍 Procurando opção '{valor}' em {len(opts10)} seletores...")
            escolhido = False
            for idx, ox in enumerate(opts10, 1):
                try:
                    lis = self.driver.find_elements(By.XPATH, ox)
                    logger.debug(f"    Seletor opção {idx}: {ox} -> {len(lis)} elementos")
                    vis = []
                    for li in lis:
                        try:
                            if li.is_displayed():
                                vis.append(li)
                        except Exception:
                            continue
                    if vis:
                        logger.info(f"  ✓ Opção '{valor}' encontrada com seletor {idx}")
                        self._scroll_center(vis[0])
                        if self._safe_js_click(vis[0]):
                            escolhido = True
                            logger.info(f"  ✓ Paginação configurada para {valor} itens")
                            break
                        try:
                            self.driver.execute_script("arguments[0].click();", vis[0])
                            escolhido = True
                            logger.info(f"  ✓ Paginação configurada para {valor} itens (JS)")
                            break
                        except Exception:
                            continue
                except Exception as e:
                    logger.debug(f"    Erro no seletor opção {idx}: {e}")
                    continue
            if not escolhido:
                logger.warning(f"⚠ Opção '{valor}' não encontrada no dropdown após testar todos os seletores.")
                logger.warning(f"   Tentou {len(opts10)} seletores diferentes")

                # Último fallback: listar todas as opções visíveis para debug
                logger.info("  🔍 Tentando fallback: procurando TODAS as opções visíveis...")
                try:
                    all_lis = self.driver.find_elements(By.XPATH, "//ul[@role='listbox']//li")
                    visible_options = []
                    for li in all_lis:
                        try:
                            if li.is_displayed():
                                text = li.text.strip()
                                visible_options.append(text)
                                logger.debug(f"    Opção encontrada: '{text}'")
                                if text == valor or valor in text:
                                    logger.info(f"  ✓ Encontrado '{valor}' via fallback!")
                                    self._scroll_center(li)
                                    if self._safe_js_click(li):
                                        escolhido = True
                                        logger.info(f"  ✓ Paginação configurada para {valor} itens (fallback)")
                                        break
                        except Exception:
                            continue

                    if not escolhido and visible_options:
                        logger.warning(f"  📋 Opções disponíveis: {visible_options}")
                except Exception as e:
                    logger.debug(f"  Erro no fallback: {e}")

                if not escolhido:
                    return False
            # Aguardar recarregar
            logger.info("  ⏳ Aguardando resultados recarregarem...")
            time.sleep(0.8)
            self._wait_results_loaded(20)
            logger.info("✅ Itens por página ajustado com sucesso!")
            return True
        except Exception as e:
            logger.error(f"Erro ajustando itens por página: {e}")
            return False

    def _find_next_page_button(self):
        try:
            pag = self.selectors.get('jt', {}).get('resultados', {}).get('paginacao', {})
            # Dentro do container
            containers = []
            for xp in pag.get('paginator_container', []) or []:
                try:
                    cs = self.driver.find_elements(By.XPATH, xp)
                    containers.extend([c for c in cs if (c.is_displayed() if hasattr(c,'is_displayed') else True)])
                except Exception:
                    continue
            # Tentar achar botão dentro do container
            for cont in containers or [self.driver]:
                for bx in pag.get('next_button', []) or []:
                    try:
                        bs = (cont.find_elements(By.XPATH, bx) if cont is not self.driver else self.driver.find_elements(By.XPATH, bx))
                        vis = []
                        for b in bs:
                            try:
                                displayed = b.is_displayed() if hasattr(b, 'is_displayed') else True
                            except Exception:
                                displayed = True
                            if not displayed:
                                continue
                            # Ignorar botões desabilitados
                            try:
                                cls = (b.get_attribute('class') or '').lower()
                                dis_attr = (b.get_attribute('disabled') or '').lower()
                                aria_dis = (b.get_attribute('aria-disabled') or '').lower()
                                if 'p-disabled' in cls or dis_attr in ('true', 'disabled', '1') or aria_dis in ('true', 'disabled', '1'):
                                    continue
                            except Exception:
                                pass
                            vis.append(b)
                        if vis:
                            return vis[0]
                    except Exception:
                        continue
                # fallback: span ícone
                for sx in pag.get('next_icon_span', []) or []:
                    try:
                        spans = (cont.find_elements(By.XPATH, sx) if cont is not self.driver else self.driver.find_elements(By.XPATH, sx))
                        vis = []
                        for s in spans:
                            try:
                                displayed = s.is_displayed() if hasattr(s, 'is_displayed') else True
                            except Exception:
                                displayed = True
                            if not displayed:
                                continue
                            vis.append(s)
                        if vis:
                            # pegar botão ancestral e checar se não está desabilitado
                            try:
                                btn = vis[0].find_element(By.XPATH, "ancestor::button[1]")
                                try:
                                    cls = (btn.get_attribute('class') or '').lower()
                                    dis_attr = (btn.get_attribute('disabled') or '').lower()
                                    aria_dis = (btn.get_attribute('aria-disabled') or '').lower()
                                    if 'p-disabled' in cls or dis_attr in ('true', 'disabled', '1') or aria_dis in ('true', 'disabled', '1'):
                                        btn = None
                                except Exception:
                                    pass
                                if btn is not None:
                                    return btn
                            except Exception:
                                pass
                    except Exception:
                        continue
            return None
        except Exception:
            return None

    def _get_pagination_info(self) -> Tuple[Optional[int], Optional[int]]:
        """Lê a paginação do UI: página atual e maior página visível.
        Retorna (current_page, last_page) ou (None, None) se não conseguir.
        """
        try:
            # Procurar containers do paginator
            containers = []
            try:
                for xp in (self.selectors.get('jt', {}).get('resultados', {}).get('paginacao', {}).get('paginator_container', []) or []):
                    try:
                        cs = self.driver.find_elements(By.XPATH, xp)
                        containers.extend([c for c in cs if (c.is_displayed() if hasattr(c,'is_displayed') else True)])
                    except Exception:
                        continue
            except Exception:
                pass
            if not containers:
                return None, None
            current = None
            last = None
            for cont in containers:
                try:
                    # Botões de página
                    btns = cont.find_elements(By.XPATH, ".//button[contains(@class,'p-paginator-page')]")
                except Exception:
                    btns = []
                for b in btns:
                    try:
                        if not (b.is_displayed() if hasattr(b,'is_displayed') else True):
                            continue
                        txt = (b.text or '').strip()
                        n = int(re.sub(r"\D+", "", txt)) if re.search(r"\d", txt) else None
                        if n is not None:
                            last = (n if last is None else max(last, n))
                            cls = (b.get_attribute('class') or '').lower()
                            if 'p-highlight' in cls:
                                current = n
                    except Exception:
                        continue
            return current, last
        except Exception:
            return None, None

    def _go_to_page_one(self, timeout: int = 25) -> bool:
        """Força o paginator para a página 1, se possível.
        Tenta clicar no botão '1' dentro do paginator; caso não exista, tenta usar o botão 'prev' até desabilitar.
        """
        try:
            # garantir visibilidade do paginator
            self._scroll_to_bottom(tries=2)
            containers = []
            try:
                for xp in (self.selectors.get('jt', {}).get('resultados', {}).get('paginacao', {}).get('paginator_container', []) or []):
                    try:
                        cs = self.driver.find_elements(By.XPATH, xp)
                        containers.extend([c for c in cs if (c.is_displayed() if hasattr(c,'is_displayed') else True)])
                    except Exception:
                        continue
            except Exception:
                pass
            if not containers:
                return False
            # Tentar botão da página '1'
            for cont in containers:
                try:
                    btns = cont.find_elements(By.XPATH, ".//button[contains(@class,'p-paginator-page') and normalize-space()='1']")
                except Exception:
                    btns = []
                for b in btns:
                    try:
                        if not (b.is_displayed() if hasattr(b,'is_displayed') else True):
                            continue
                        cls = (b.get_attribute('class') or '').lower()
                        dis_attr = (b.get_attribute('disabled') or '').lower()
                        aria_dis = (b.get_attribute('aria-disabled') or '').lower()
                        if 'p-disabled' in cls or dis_attr in ('true', 'disabled', '1') or aria_dis in ('true', 'disabled', '1'):
                            continue
                        self._scroll_center(b)
                        self._hover(b)
                        if not self._safe_js_click(b):
                            try:
                                self.driver.execute_script("arguments[0].click();", b)
                            except Exception:
                                pass
                        ok = self._wait_results_loaded(timeout)
                        if ok:
                            logger.info("Paginator reposicionado na página 1.")
                            return True
                    except Exception:
                        continue
            # Fallback: tentar botão 'prev' até desabilitar
            for cont in containers:
                try:
                    prev = cont.find_elements(By.XPATH, ".//button[contains(@class,'p-paginator-prev')]")
                except Exception:
                    prev = []
                if not prev:
                    continue
                pbtn = prev[0]
                for _ in range(12):  # alguns passos para garantir voltar ao início
                    try:
                        cls = (pbtn.get_attribute('class') or '').lower()
                        dis_attr = (pbtn.get_attribute('disabled') or '').lower()
                        aria_dis = (pbtn.get_attribute('aria-disabled') or '').lower()
                        if 'p-disabled' in cls or dis_attr in ('true', 'disabled', '1') or aria_dis in ('true', 'disabled', '1'):
                            break
                        self._scroll_center(pbtn)
                        self._hover(pbtn)
                        if not self._safe_js_click(pbtn):
                            try:
                                self.driver.execute_script("arguments[0].click();", pbtn)
                            except Exception:
                                pass
                        self._wait_results_loaded(10)
                    except Exception:
                        break
                # checar se na página 1 agora
                cur, _ = self._get_pagination_info()
                if cur == 1:
                    logger.info("Paginator reposicionado via 'prev' para página 1.")
                    return True
            return False
        except Exception:
            return False

    def _go_to_next_page(self) -> bool:
        try:
            # rolar ao fundo (garantir paginator visível)
            self._scroll_to_bottom(tries=3)
            # Guarda de segurança: não passar da página 10 pelo estado do UI
            try:
                cur_pg, last_pg = self._get_pagination_info()
                if cur_pg is not None and cur_pg >= 10:
                    logger.info(f"Página atual no UI é {cur_pg} (>= 10). Não avançará.")
                    return False
                # Também evitar avançar se já estivermos na última página visível
                if cur_pg is not None and last_pg is not None and cur_pg >= last_pg:
                    logger.info(f"Última página do paginator alcançada ({cur_pg}/{last_pg}). Não avançará.")
                    return False
            except Exception:
                pass
            btn = self._find_next_page_button()
            if not btn:
                logger.info("Botão próxima página não encontrado (talvez última página).")
                return False
            # Verificar explicitamente se o botão está desabilitado
            try:
                cls = (btn.get_attribute('class') or '').lower()
                dis_attr = (btn.get_attribute('disabled') or '').lower()
                aria_dis = (btn.get_attribute('aria-disabled') or '').lower()
                if 'p-disabled' in cls or dis_attr in ('true', 'disabled', '1') or aria_dis in ('true', 'disabled', '1'):
                    logger.info("Botão 'Próxima' está desabilitado. Última página alcançada.")
                    return False
            except Exception:
                pass
            self._scroll_center(btn)
            self._hover(btn)
            time.sleep(0.05)
            if not self._safe_js_click(btn):
                try:
                    self.driver.execute_script("arguments[0].dispatchEvent(new MouseEvent('mousedown',{bubbles:true,buttons:1}));", btn)
                    self.driver.execute_script("arguments[0].dispatchEvent(new MouseEvent('mouseup',{bubbles:true,buttons:1}));", btn)
                    self.driver.execute_script("arguments[0].click();", btn)
                except Exception:
                    pass
            # aguardar carregar nova página
            ok = self._wait_results_loaded(25)
            if ok:
                logger.info("Avançou para a próxima página.")
            return ok
        except Exception as e:
            logger.error(f"Erro ao avançar página: {e}")
            return False

    def _get_center_y(self, el) -> Optional[float]:
        try:
            r = el.rect
            return float(r.get('y', 0.0)) + float(r.get('height', 0.0)) / 2.0
        except Exception:
            return None

    def _focus_card(self, card):
        # Tenta dar foco/seleção ao cartão clicando numa área segura
        try:
            # preferir um bloco de texto do cabeçalho
            cand = None
            try:
                cand = card.find_elements(By.XPATH, ".//div[contains(@class,'doc-texto')]")
                cand = cand[0] if cand else None
            except Exception:
                cand = None
            if cand is not None:
                self._scroll_center(cand)
                self._hover(cand)
                try:
                    cand.click()
                except Exception:
                    try:
                        self.driver.execute_script("arguments[0].click();", cand)
                    except Exception:
                        pass
        except Exception:
            pass

    def _pick_nearest_element(self, card, elements: Iterable) -> Optional[object]:
        try:
            cy = self._get_center_y(card)
            if cy is None:
                return None
            best = None
            best_d = None
            for b in elements:
                try:
                    if not b.is_displayed():
                        continue
                except Exception:
                    continue
                by = self._get_center_y(b)
                if by is None:
                    continue
                d = abs(by - cy)
                if best is None or d < best_d:
                    best = b
                    best_d = d
            return best
        except Exception:
            return None

    def _find_copy_element(self, card):
        sel = self.selectors.get('jt', {}).get('resultados', {})

        # DEBUG: Verifica se é modo debug
        debug_mode = os.environ.get("JT_DEBUG_COPY", "0") == "1"

        # 1) PRIORIDADE: Tentar seletores específicos de EMENTA primeiro
        rel_paths = sel.get('copiar_ementa', [])
        if debug_mode:
            logger.debug(f"  [1] Tentando seletores específicos de ementa (total: {len(rel_paths)})")
        for idx, xp in enumerate(rel_paths):
            try:
                if debug_mode:
                    logger.debug(f"      [{idx+1}] XPath: {xp}")
                btns = card.find_elements(By.XPATH, xp)
                if debug_mode:
                    logger.debug(f"          Encontrados: {len(btns)} elementos")
                vis = []
                for b in btns:
                    try:
                        if b.is_displayed():
                            vis.append(b)
                    except Exception:
                        continue
                if vis:
                    if debug_mode:
                        logger.debug(f"          ✓ Retornando elemento visível (específico de ementa)")
                    return vis[0]
            except Exception as e:
                if debug_mode:
                    logger.debug(f"          Erro: {e}")
                continue

        # 2) Fallback: span genérico (só se não encontrou o específico)
        if debug_mode:
            logger.debug(f"  [2] Fallback: Tentando copiar_span genérico")
        span_rel = sel.get('copiar_span')
        if span_rel:
            try:
                spans = card.find_elements(By.XPATH, span_rel)
                if debug_mode:
                    logger.debug(f"      Encontrados: {len(spans)} spans")
                vis = []
                for s in spans:
                    try:
                        if s.is_displayed():
                            vis.append(s)
                    except Exception:
                        continue
                if vis:
                    if debug_mode:
                        logger.debug(f"      ✓ Retornando span visível")
                    return vis[0]
            except Exception as e:
                if debug_mode:
                    logger.debug(f"      Erro: {e}")
                pass

        # 3) Tentar botão relativo ao cartão (segunda tentativa)
        for xp in rel_paths:
            try:
                btns = card.find_elements(By.XPATH, xp)
                vis = []
                for b in btns:
                    try:
                        if b.is_displayed():
                            vis.append(b)
                    except Exception:
                        continue
                if vis:
                    return vis[0]
            except Exception:
                continue
        # 3) Subir para contêiner pai e tentar novamente relativo
        try:
            parents = card.find_elements(By.XPATH, "ancestor::*[contains(@class,'documento') or contains(@class,'doc')][1]")
            if parents:
                cont = parents[0]
                for xp in rel_paths:
                    try:
                        btns = cont.find_elements(By.XPATH, xp)
                        vis = [b for b in btns if (b.is_displayed() if hasattr(b,'is_displayed') else True)]
                        if vis:
                            return vis[0]
                    except Exception:
                        continue
        except Exception:
            pass
        # 4) Global: span 'content_copy' mais próximo
        span_glob = sel.get('copiar_span_global')
        glob_spans = []
        if span_glob:
            try:
                glob_spans = self.driver.find_elements(By.XPATH, span_glob)
            except Exception:
                glob_spans = []
        if glob_spans:
            el = self._pick_nearest_element(card, glob_spans)
            if el is not None:
                return el
        # 5) Global: botão mais próximo
        glob_paths = sel.get('copiar_ementa_global', [])
        glob_btns = []
        for gx in glob_paths:
            try:
                bs = self.driver.find_elements(By.XPATH, gx)
                glob_btns.extend(bs)
            except Exception:
                continue
        if glob_btns:
            return self._pick_nearest_element(card, glob_btns)
        return None

    def _safe_js_click(self, el):
        # Preferir clique nativo (conta como gesto do usuário p/ clipboard)
        try:
            self._scroll_center(el)
            try:
                el.click()
                return True
            except Exception:
                pass
            # Fallback: forçar foco e clique via JS
            try:
                self.driver.execute_script("arguments[0].focus();", el)
            except Exception:
                pass
            self.driver.execute_script("arguments[0].click();", el)
            return True
        except Exception:
            return False

    def _wait_results_loaded(self, timeout=30) -> bool:
        inicio = time.time()
        last_vis_count = -1
        stable = 0
        sel = self.selectors.get('jt', {}).get('resultados', {})
        cartao_paths = sel.get('cartao_root', [])
        while time.time() - inicio < timeout:
            vis = []
            for xp in cartao_paths:
                try:
                    elems = self.driver.find_elements(By.XPATH, xp)
                    for e in elems:
                        try:
                            if e.is_displayed():
                                vis.append(e)
                        except Exception:
                            continue
                except Exception:
                    continue
            count = len(vis)
            if count == last_vis_count:
                stable += 1
            else:
                stable = 0
                last_vis_count = count
            # Checa spinners
            try:
                spinners = self.driver.find_elements(By.XPATH, "//*[contains(@class,'spinner') or contains(@class,'progress') or contains(@class,'CircularProgress')]")
            except Exception:
                spinners = []
            if stable >= 2 and (count > 0 or (time.time() - inicio) > 5) and not spinners:
                logger.info(f"Resultados prontos. Cartões visíveis: {count}")
                return True
            time.sleep(0.6)
        logger.warning("Timeout aguardando resultados.")
        return False

    def _expandir_mais_tribunais(self, max_clicks=5) -> bool:
        """Clica em 'Mais...' para revelar mais tribunais (ex.: TRT24)."""
        filtros = self.selectors.get('jt', {}).get('filtros', {})
        paths = filtros.get('mais_tribunais', [])
        clicks = 0
        for _ in range(max(1, max_clicks)):
            found_btn = None
            for xp in paths:
                try:
                    els = self.driver.find_elements(By.XPATH, xp)
                    vis = []
                    for e in els:
                        try:
                            if e.is_displayed():
                                vis.append(e)
                        except Exception:
                            continue
                    if vis:
                        found_btn = vis[-1]
                        break
                except Exception:
                    continue
            if not found_btn:
                break
            if self._safe_js_click(found_btn):
                clicks += 1
                time.sleep(0.5)
            else:
                break
        if clicks:
            logger.info(f"Expandido 'Mais...' {clicks} vez(es).")
        return clicks > 0

    def _try_click_pesquisar(self):
        cfg = self.selectors.get('jt', {})
        geral = cfg.get('geral', {})
        paths = geral.get('botao_pesquisar', [])
        for xp in paths:
            try:
                btns = self.driver.find_elements(By.XPATH, xp)
                vis = []
                for b in btns:
                    try:
                        if b.is_displayed():
                            vis.append(b)
                    except Exception:
                        continue
                if vis:
                    ok = self._safe_js_click(vis[0])
                    logger.info(f"Clique em 'Pesquisar': {'ok' if ok else 'falha'}")
                    time.sleep(0.8)
                    return
            except Exception:
                continue
        logger.info("Botão 'Pesquisar' não localizado (ignorado).")

    # ---------- Fluxo Principal ----------
    def test_extrair_jt(self):
        cfg = self.selectors.get('jt', {})
        url = cfg.get('url')
        if not url:
            self.fail("URL do JT não encontrada em selectors_jt.json")

        logger.info("Abrindo portal JT...")
        self.open(url)
        self.wait_for_element("body", timeout=30)
        time.sleep(1.2)

        filtros = cfg.get('filtros', {})
        xp_acordaos = filtros.get('acordaos')
        xp_com_ementa = filtros.get('com_ementa')
        xp_tribunal_tmpl = filtros.get('tribunal_label')

        # Helpers locais para seleção estável
        def verificar_filtro_marcado(elemento, debug: bool = False) -> bool:
            """Verifica se um filtro/checkbox está marcado"""
            try:
                # Coleta informações para debug
                if debug:
                    classes = elemento.get_attribute('class') or ''
                    aria_checked = elemento.get_attribute('aria-checked')
                    aria_selected = elemento.get_attribute('aria-selected')
                    logger.info(f"    [DEBUG] Classes: {classes}")
                    logger.info(f"    [DEBUG] aria-checked: {aria_checked}")
                    logger.info(f"    [DEBUG] aria-selected: {aria_selected}")

                # Verifica por atributos comuns de estado marcado
                aria_checked = elemento.get_attribute('aria-checked')
                if aria_checked == 'true':
                    return True

                aria_selected = elemento.get_attribute('aria-selected')
                if aria_selected == 'true':
                    return True

                # Verifica classes que indicam seleção (PrimeNG usa p-highlight)
                classes = elemento.get_attribute('class') or ''
                marcadores = ['selected', 'checked', 'active', 'p-highlight', 'p-checked', 'selecionado']
                if any(x in classes for x in marcadores):
                    return True

                # Verifica elemento pai (às vezes a classe está no pai)
                try:
                    pai = elemento.find_element(By.XPATH, "..")
                    classes_pai = pai.get_attribute('class') or ''
                    if any(x in classes_pai for x in marcadores):
                        return True
                except:
                    pass

                # Verifica input checkbox dentro do elemento
                try:
                    checkbox = elemento.find_element(By.XPATH, ".//input[@type='checkbox']")
                    if checkbox.is_selected():
                        return True
                except:
                    pass

                # Verifica checkbox no elemento pai
                try:
                    pai = elemento.find_element(By.XPATH, "..")
                    checkbox_pai = pai.find_element(By.XPATH, ".//input[@type='checkbox']")
                    if checkbox_pai.is_selected():
                        return True
                except:
                    pass

                return False
            except Exception as e:
                if debug:
                    logger.warning(f"    [DEBUG] Erro verificando: {e}")
                return False

        def aplicar_filtros_basicos():
            logger.info("📋 Aplicando filtros básicos...")

            # Filtro Acórdãos
            if xp_acordaos:
                logger.info("  🔍 Aplicando filtro 'Acórdãos'...")
                for tentativa in range(3):
                    try:
                        btns = self.driver.find_elements(By.XPATH, xp_acordaos)
                        if btns:
                            btn = btns[0]
                            # Verifica se já está marcado
                            if verificar_filtro_marcado(btn):
                                logger.info("    ✓ 'Acórdãos' já marcado")
                                break

                            self._scroll_center(btn)
                            self._safe_js_click(btn)
                            time.sleep(0.8)

                            # Verifica se marcou
                            if verificar_filtro_marcado(btn):
                                logger.info("    ✓ 'Acórdãos' marcado com sucesso")
                                break
                            else:
                                logger.warning(f"    ⚠ Tentativa {tentativa+1}: 'Acórdãos' não marcou")
                    except Exception as e:
                        logger.debug(f"    Erro na tentativa {tentativa+1}: {e}")
                        time.sleep(0.5)

            # Filtro Com Ementa
            if xp_com_ementa:
                logger.info("  🔍 Aplicando filtro 'Com Ementa'...")
                for tentativa in range(3):
                    try:
                        btns_ce = self.driver.find_elements(By.XPATH, xp_com_ementa)
                        if btns_ce:
                            btn_ce = btns_ce[0]
                            # Verifica se já está marcado
                            if verificar_filtro_marcado(btn_ce):
                                logger.info("    ✓ 'Com Ementa' já marcado")
                                break

                            self._scroll_center(btn_ce)
                            self._safe_js_click(btn_ce)
                            time.sleep(0.8)

                            # Verifica se marcou
                            if verificar_filtro_marcado(btn_ce):
                                logger.info("    ✓ 'Com Ementa' marcado com sucesso")
                                break
                            else:
                                logger.warning(f"    ⚠ Tentativa {tentativa+1}: 'Com Ementa' não marcou")
                    except Exception as e:
                        logger.debug(f"    Erro na tentativa {tentativa+1}: {e}")
                        time.sleep(0.5)

            logger.info("✅ Filtros básicos aplicados")

        # REMOVIDO: Não é necessário clicar em Pesquisar - os filtros já atualizam automaticamente

        def limpar_todos_filtros_turma_forcado():
            """
            Função de fallback para remover TODOS os filtros de turma de forma mais agressiva.
            Usa JavaScript direto para forçar remoção quando métodos normais falham.

            Retorna:
                bool: True se conseguiu limpar, False caso contrário
            """
            try:
                logger.info("  🔧 Tentando limpeza forçada de filtros de turma (fallback)...")

                # Script JS para remover todos os chips que contêm 'Turma' mas não 'TRT'
                js_script = """
                var chips = document.querySelectorAll('div.p-chip');
                var removidos = 0;
                chips.forEach(function(chip) {
                    var texto = chip.innerText || chip.textContent;
                    if (texto && texto.includes('Turma') && !texto.includes('TRT')) {
                        var removeBtn = chip.querySelector('span.pi-times-circle, span.pi-chip-remove-icon');
                        if (removeBtn) {
                            removeBtn.click();
                            removidos++;
                        }
                    }
                });
                return removidos;
                """

                removidos = self.driver.execute_script(js_script)
                logger.info(f"    ✓ Limpeza forçada removeu {removidos} chip(s) via JavaScript")

                # Aguarda UI processar
                time.sleep(1.5)

                return removidos > 0

            except Exception as e:
                logger.error(f"    ❌ Erro na limpeza forçada: {e}")
                return False

        def desmarcar_turma_atual():
            """Desmarca TODAS as turmas atualmente selecionadas clicando no X (somente turmas, não TRTs)

            Retorna:
                dict: {'sucesso': bool, 'total_removidos': int, 'detalhes': list}
            """
            resultado = {
                'sucesso': False,
                'total_removidos': 0,
                'detalhes': [],
                'erros': []
            }

            try:
                logger.info("  🧹 Iniciando remoção de filtros de turmas...")

                # XPaths alternativos para diferentes estruturas de chips
                xpaths_remove = [
                    # XPath original - chips com 'ª Turma'
                    (
                        "//div[contains(@class,'p-chip')]"
                        "[.//div[contains(@class,'p-chip-text') and contains(normalize-space(),'ª Turma')]]"
                        "//span[contains(@class,'pi-chip-remove-icon') and contains(@class,'pi-times-circle')]"
                    ),
                    # XPath alternativo 1 - mais genérico para turmas
                    (
                        "//div[contains(@class,'p-chip')]"
                        "[contains(.,'Turma')]"
                        "//span[contains(@class,'pi-times-circle')]"
                    ),
                    # XPath alternativo 2 - busca por padrão de número + ª
                    (
                        "//div[contains(@class,'p-chip')]"
                        "[.//div[contains(@class,'p-chip-text') and contains(text(),'ª')]]"
                        "[not(contains(.,'TRT'))]"
                        "//span[contains(@class,'pi-chip-remove-icon')]"
                    ),
                ]

                max_ciclos = 15  # Aumentado de 10 para 15
                turmas_removidas_nomes = []

                for ciclo in range(max_ciclos):
                    encontrou_algum = False

                    # Tenta todos os XPaths
                    for idx, xp_remove in enumerate(xpaths_remove):
                        try:
                            removes = self.driver.find_elements(By.XPATH, xp_remove)
                            logger.debug(f"    [Ciclo {ciclo+1}] XPath {idx+1} encontrou {len(removes)} elementos")
                        except Exception as e:
                            logger.debug(f"    [Ciclo {ciclo+1}] Erro ao buscar com XPath {idx+1}: {e}")
                            removes = []

                        # Filtra apenas elementos visíveis
                        visiveis = []
                        for rem in removes:
                            try:
                                if rem.is_displayed():
                                    visiveis.append(rem)
                            except Exception:
                                continue

                        if not visiveis:
                            continue

                        encontrou_algum = True
                        logger.info(f"    [Ciclo {ciclo+1}] Encontrados {len(visiveis)} chip(s) de turma para remover (XPath {idx+1})")

                        # Clica em todos os X visíveis
                        for rem in visiveis:
                            try:
                                # Tenta extrair o nome da turma antes de remover
                                turma_nome = "Desconhecida"
                                try:
                                    chip_parent = rem.find_element(By.XPATH, "../..")
                                    turma_nome = chip_parent.text.strip()
                                except Exception:
                                    pass

                                logger.debug(f"      Removendo turma: {turma_nome}")

                                # Scroll e clique
                                self._scroll_center(rem)
                                time.sleep(0.2)  # Pequena pausa para estabilizar

                                click_ok = self._safe_js_click(rem)

                                if not click_ok:
                                    # Fallback: JS direto
                                    logger.debug(f"      Tentando JS click direto para {turma_nome}")
                                    try:
                                        self.driver.execute_script("arguments[0].click();", rem)
                                        click_ok = True
                                    except Exception as e_js:
                                        logger.warning(f"      ✗ Falha ao remover {turma_nome}: {e_js}")
                                        resultado['erros'].append(f"Falha ao clicar em {turma_nome}: {str(e_js)}")
                                        continue

                                if click_ok:
                                    resultado['total_removidos'] += 1
                                    turmas_removidas_nomes.append(turma_nome)
                                    resultado['detalhes'].append(f"Removida: {turma_nome}")
                                    logger.debug(f"      ✓ Removida: {turma_nome}")

                                    # Aguarda UI atualizar (aumentado de 0.4s para 0.8s)
                                    time.sleep(0.8)

                            except Exception as e:
                                logger.debug(f"      Erro removendo chip: {e}")
                                resultado['erros'].append(f"Erro ao processar chip: {str(e)}")
                                continue

                        # Se removeu chips com este XPath, pausa antes de tentar próximo
                        if visiveis:
                            time.sleep(0.5)

                    # Se não encontrou mais nada em nenhum XPath, encerra
                    if not encontrou_algum:
                        logger.info(f"    [Ciclo {ciclo+1}] Nenhum chip de turma encontrado. Finalizando limpeza.")
                        break

                    # Pausa entre ciclos
                    time.sleep(0.6)

                # Validação final: verifica se ainda há chips de turma
                time.sleep(0.5)
                chips_restantes = []
                for xp in xpaths_remove:
                    try:
                        chips = self.driver.find_elements(By.XPATH, xp)
                        chips_vis = [c for c in chips if c.is_displayed()]
                        chips_restantes.extend(chips_vis)
                    except Exception:
                        pass

                if chips_restantes:
                    logger.warning(f"    ⚠ ATENÇÃO: Ainda restam {len(chips_restantes)} chip(s) de turma após limpeza!")

                    # Tenta limpeza forçada como fallback
                    logger.info(f"    🔄 Ativando modo de limpeza forçada...")
                    if limpar_todos_filtros_turma_forcado():
                        # Valida novamente
                        time.sleep(0.5)
                        chips_apos_forcado = []
                        for xp in xpaths_remove:
                            try:
                                chips = self.driver.find_elements(By.XPATH, xp)
                                chips_vis = [c for c in chips if c.is_displayed()]
                                chips_apos_forcado.extend(chips_vis)
                            except Exception:
                                pass

                        if not chips_apos_forcado:
                            logger.info(f"    ✅ Limpeza forçada bem-sucedida!")
                            resultado['sucesso'] = True
                        else:
                            logger.error(f"    ❌ Ainda restam {len(chips_apos_forcado)} chips após limpeza forçada")
                            resultado['sucesso'] = False
                            resultado['erros'].append(f"{len(chips_apos_forcado)} chips não puderam ser removidos mesmo com limpeza forçada")
                    else:
                        resultado['sucesso'] = False
                        resultado['erros'].append(f"{len(chips_restantes)} chips não puderam ser removidos")
                else:
                    logger.info(f"    ✓ Validação: Nenhum chip de turma remanescente")
                    resultado['sucesso'] = True

                # Sumário
                if resultado['total_removidos'] > 0:
                    logger.info(f"    ✅ Total de {resultado['total_removidos']} turma(s) desmarcada(s): {', '.join(turmas_removidas_nomes)}")
                else:
                    logger.info(f"    ℹ Nenhuma turma estava marcada")

                if resultado['erros']:
                    logger.warning(f"    ⚠ {len(resultado['erros'])} erro(s) durante a remoção")

                return resultado

            except Exception as e:
                logger.error(f"    ❌ Erro crítico ao desmarcar turmas: {e}")
                logger.debug(traceback.format_exc())
                resultado['erros'].append(f"Erro crítico: {str(e)}")
                return resultado

        def selecionar_trt(sigla: str) -> bool:
            logger.info(f"  🏛️ Selecionando tribunal {sigla}...")
            try:
                if not xp_tribunal_tmpl:
                    return False
                xp = None
                try:
                    xp = xp_tribunal_tmpl.format(sigla)
                except Exception:
                    xp = None
                if not xp:
                    return False

                # Tenta até 3 vezes
                for tentativa in range(3):
                    itens = self.driver.find_elements(By.XPATH, xp)

                    # Se não encontrar TRT24, expande lista
                    if not itens and sigla == "TRT24":
                        try:
                            logger.info("    📂 Expandindo lista para encontrar TRT24...")
                            self._expandir_mais_tribunais(max_clicks=5)
                            time.sleep(0.5)
                        except Exception:
                            pass
                        itens = self.driver.find_elements(By.XPATH, xp)

                    if itens:
                        # IMPORTANTE: O seletor retorna o <span>, mas precisamos do <div> pai que é clicável
                        span_element = itens[0]

                        # Pega o elemento pai (div.filtro-item) que é o elemento interativo
                        try:
                            item = span_element.find_element(By.XPATH, "..")
                            item_tag = item.tag_name
                            item_classes = item.get_attribute('class') or ''
                            logger.info(f"    🔍 Tentativa {tentativa+1}: Elemento pai <{item_tag}> (classes: {item_classes[:50]}...)")
                        except:
                            # Fallback: usa o próprio span se não conseguir pegar o pai
                            item = span_element
                            logger.info(f"    🔍 Tentativa {tentativa+1}: Usando elemento span")

                        # Verifica se já está marcado (com debug na primeira tentativa)
                        usa_debug = (tentativa == 0)
                        if verificar_filtro_marcado(item, debug=usa_debug):
                            logger.info(f"    ✓ {sigla} já marcado")
                            return True

                        # Clica para marcar
                        logger.info(f"    🖱️ Clicando em {sigla}...")
                        self._scroll_center(item)
                        click_sucesso = self._safe_js_click(item)

                        if click_sucesso:
                            # Para TRT, quando clica o elemento MUDA DE LUGAR (vai para área de filtros aplicados)
                            time.sleep(1.5)  # Aguarda UI atualizar completamente

                            # Verifica se elemento ainda existe no local original
                            itens_apos_click = self.driver.find_elements(By.XPATH, xp)

                            if not itens_apos_click:
                                # Elemento sumiu do local original = foi movido para filtros aplicados = SUCESSO!
                                logger.info(f"    ✓ {sigla} selecionado com sucesso (movido para filtros aplicados)")
                                return True
                            else:
                                # Elemento ainda está no local original - pode ser que não clicou corretamente
                                logger.warning(f"    ⚠ Tentativa {tentativa+1}: Elemento {sigla} não foi movido após clique")
                        else:
                            logger.warning(f"    ⚠ Tentativa {tentativa+1}: Falha ao executar clique em {sigla}")
                    else:
                        logger.warning(f"    ⚠ Tentativa {tentativa+1}: {sigla} não encontrado no DOM")

                    # Pausa antes de tentar novamente
                    if tentativa < 2:
                        time.sleep(0.7)

                logger.warning(f"  ❌ Falha ao selecionar {sigla} após 3 tentativas")
                return False
            except Exception as e:
                logger.error(f"  ❌ Erro selecionando {sigla}: {e}")
                logger.debug(traceback.format_exc())
                return False

        def expandir_lista_turmas(max_tentativas: int = 3) -> bool:
            """Expande a lista de turmas clicando em 'Mais...'"""
            mais_paths = filtros.get('mais_turmas')
            if not mais_paths:
                return False

            # Se mais_paths é string, converte para lista
            if isinstance(mais_paths, str):
                mais_paths = [mais_paths]

            expansoes_realizadas = 0
            for _ in range(max_tentativas):
                encontrou = False
                for xpath in mais_paths:
                    try:
                        botoes = self.driver.find_elements(By.XPATH, xpath)
                        for btn in botoes:
                            if btn.is_displayed():
                                self._scroll_center(btn)
                                time.sleep(0.3)  # Pausa antes do clique
                                self._safe_js_click(btn)
                                expansoes_realizadas += 1
                                logger.info(f"✓ Lista de turmas expandida (expansão #{expansoes_realizadas})")

                                # CRÍTICO: Aguarda UI carregar novos elementos
                                time.sleep(1.2)  # Aumentado de 0.4s para 1.2s

                                # Pequeno scroll para forçar renderização
                                try:
                                    self._scroll_by(0.1)
                                    time.sleep(0.3)
                                except Exception:
                                    pass

                                encontrou = True
                                break
                        if encontrou:
                            break
                    except Exception:
                        continue
                if not encontrou:
                    break

            if expansoes_realizadas > 0:
                logger.info(f"  ✓ Total de {expansoes_realizadas} expansão(ões) realizadas. Aguardando estabilização...")
                time.sleep(0.8)  # Aguarda final para garantir

            return expansoes_realizadas > 0

        def selecionar_turma(nome: str, tentar_expandir: bool = True) -> bool:
            logger.info(f"  📋 Selecionando turma '{nome}'...")

            # Primeiro desmarca a turma atual
            resultado_desmarcacao = desmarcar_turma_atual()

            # Verifica se houve problemas na desmarcação
            if resultado_desmarcacao['erros'] and not resultado_desmarcacao['sucesso']:
                logger.warning(f"  ⚠ Problemas ao desmarcar turmas anteriores: {len(resultado_desmarcacao['erros'])} erro(s)")
                # Aguarda mais tempo para estabilizar UI
                time.sleep(1.0)
            else:
                time.sleep(0.3)

            # Helper local: variantes do rótulo (com e sem zero à esquerda)
            def _variantes_turma_label_local(txt: str):
                try:
                    s = re.sub(r"\s+", " ", (txt or "").strip())
                    variantes = [s]
                    m = re.match(r"^(\d+)\s*ª\s*Turma$", s, flags=re.IGNORECASE)
                    if m:
                        num = m.group(1)
                        no_zero = str(int(num))
                        zero2 = str(num).zfill(2)
                        for v in (f"{no_zero}ª Turma", f"{zero2}ª Turma"):
                            if v not in variantes:
                                variantes.append(v)
                    # único mantendo ordem
                    seen = set()
                    uniq = []
                    for v in variantes:
                        if v not in seen:
                            uniq.append(v)
                            seen.add(v)
                    return uniq
                except Exception:
                    return [txt]

            labels = _variantes_turma_label_local(nome)

            xp_turma_tmpl = filtros.get('turma_label') or filtros.get('orgao_julgante_turma_label')
            candidatos = []
            if xp_turma_tmpl:
                for lab in labels:
                    try:
                        candidatos.append(xp_turma_tmpl.format(lab))
                    except Exception:
                        pass
            # Fallbacks por texto - seletores mais específicos
            for lab in labels:
                candidatos.extend([
                    f"//div[contains(@class,'filtro-item')]//span[contains(@class,'nome-item') and normalize-space()='{lab}']",
                    f"//div[contains(@class,'filtro-item')]//span[contains(normalize-space(),'{lab}')]",
                    f"//span[contains(@class,'nome-item') and contains(normalize-space(),'{lab}')]",
                ])

            # Tenta até 3 vezes com verificação
            for tentativa in range(3):
                for xp in candidatos:
                    try:
                        itens = self.driver.find_elements(By.XPATH, xp)
                        vis = [it for it in itens if (it.is_displayed() if hasattr(it,'is_displayed') else True)]
                        if vis:
                            span_el = vis[0]

                            # Pega o elemento pai (div.filtro-item) que é o elemento interativo
                            try:
                                el = span_el.find_element(By.XPATH, "..")
                            except:
                                el = span_el

                            # Verifica se já está marcado
                            if verificar_filtro_marcado(el):
                                logger.info(f"    ✓ Turma '{nome}' já marcada")
                                return True

                            # Tenta clicar
                            self._scroll_center(el)
                            self._hover(el)
                            if self._safe_js_click(el):
                                time.sleep(0.8)

                                # Busca novamente e pega o pai
                                itens_check = self.driver.find_elements(By.XPATH, xp)
                                if itens_check:
                                    try:
                                        el_check = itens_check[0].find_element(By.XPATH, "..")
                                    except:
                                        el_check = itens_check[0]

                                    # Verifica se marcou
                                    if verificar_filtro_marcado(el_check):
                                        logger.info(f"    ✓ Turma '{nome}' marcada com sucesso")
                                        return True
                                    else:
                                        logger.warning(f"    ⚠ Tentativa {tentativa+1}: Turma '{nome}' não marcou")
                            else:
                                # Tenta JS click como fallback
                                try:
                                    self.driver.execute_script("arguments[0].click();", el)
                                    time.sleep(0.8)

                                    # Busca novamente e pega o pai
                                    itens_check = self.driver.find_elements(By.XPATH, xp)
                                    if itens_check:
                                        try:
                                            el_check = itens_check[0].find_element(By.XPATH, "..")
                                        except:
                                            el_check = itens_check[0]

                                        if verificar_filtro_marcado(el_check):
                                            logger.info(f"    ✓ Turma '{nome}' marcada (JS) com sucesso")
                                            return True
                                        else:
                                            logger.warning(f"    ⚠ Tentativa {tentativa+1} (JS): Turma '{nome}' não marcou")
                                except Exception:
                                    pass
                    except Exception as e:
                        logger.debug(f"    Erro na tentativa {tentativa+1} com xpath: {e}")
                        continue

                # Pequena pausa antes de tentar novamente
                if tentativa < 2:
                    time.sleep(0.5)

            # Se não encontrou após 3 tentativas, tenta expandir a lista
            if tentar_expandir:
                logger.info(f"  ⚠ Turma '{nome}' não encontrada. Tentando expandir lista...")
                if expandir_lista_turmas():
                    logger.info(f"  🔍 Buscando '{nome}' após expansão...")

                    # DIAGNÓSTICO: Lista todas as turmas disponíveis após expansão
                    try:
                        xp_debug = "//div[contains(@class,'filtro-item')]//span[contains(@class,'nome-item')]"
                        todos_items = self.driver.find_elements(By.XPATH, xp_debug)
                        turmas_visiveis = []
                        for item in todos_items:
                            try:
                                if item.is_displayed():
                                    texto = item.text.strip()
                                    if 'Turma' in texto and 'TRT' not in texto:
                                        turmas_visiveis.append(texto)
                            except Exception:
                                pass
                        if turmas_visiveis:
                            logger.info(f"    📋 Turmas disponíveis após expansão: {', '.join(turmas_visiveis)}")
                        else:
                            logger.warning(f"    ⚠ Nenhuma turma visível após expansão (pode indicar problema no XPath)")
                    except Exception as e:
                        logger.debug(f"    Erro ao listar turmas: {e}")

                    # ADICIONA: XPaths ainda mais genéricos para busca pós-expansão
                    candidatos_expandidos = candidatos.copy()

                    # Adiciona XPaths super genéricos para capturar elementos que podem ter estrutura diferente
                    for lab in labels:
                        candidatos_expandidos.extend([
                            # Busca em QUALQUER div/span que contenha o texto exato
                            f"//*[contains(@class,'filtro') and contains(text(),'{lab}')]",
                            f"//*[contains(@class,'item') and contains(text(),'{lab}')]",
                            # Busca por texto direto (mais permissivo)
                            f"//*[contains(text(),'{lab}')]",
                        ])

                    # Scroll para baixo para garantir que novos elementos estejam na viewport
                    try:
                        logger.debug(f"    Scrollando para revelar novos elementos...")
                        self._scroll_by(0.5)
                        time.sleep(0.5)
                        self._scroll_by(-0.3)  # Volta um pouco
                        time.sleep(0.3)
                    except Exception:
                        pass

                    # Tenta novamente após expandir (mais 5 tentativas ao invés de 3)
                    for tentativa in range(5):
                        logger.debug(f"    Tentativa {tentativa+1}/5 após expansão...")

                        for idx, xp in enumerate(candidatos_expandidos):
                            try:
                                itens = self.driver.find_elements(By.XPATH, xp)
                                vis = [it for it in itens if (it.is_displayed() if hasattr(it,'is_displayed') else True)]

                                if vis:
                                    logger.debug(f"      XPath {idx+1} encontrou {len(vis)} elemento(s)")
                                    span_el = vis[0]

                                    # Pega o elemento pai
                                    try:
                                        el = span_el.find_element(By.XPATH, "..")
                                    except:
                                        el = span_el

                                    if verificar_filtro_marcado(el):
                                        logger.info(f"    ✓ Turma '{nome}' já marcada (após expandir)")
                                        return True

                                    # Validação: elemento está visível e habilitado?
                                    try:
                                        is_displayed = el.is_displayed()
                                        is_enabled = el.is_enabled()
                                        logger.debug(f"      Estado: displayed={is_displayed}, enabled={is_enabled}")
                                        if not is_displayed:
                                            logger.debug(f"      ⚠ Elemento não visível, pulando...")
                                            continue
                                    except Exception as e:
                                        logger.debug(f"      Erro verificar estado: {e}")

                                    # ESTRATÉGIA MÚLTIPLA DE CLIQUE
                                    logger.debug(f"      🎯 Testando 5 métodos de clique para '{nome}'...")

                                    # Scroll + Hover preparatório
                                    self._scroll_center(el)
                                    time.sleep(0.5)
                                    self._hover(el)
                                    time.sleep(0.4)

                                    # Método 1: _safe_js_click
                                    logger.debug(f"      [1/5] _safe_js_click...")
                                    if self._safe_js_click(el):
                                        time.sleep(1.2)
                                        itens_check = self.driver.find_elements(By.XPATH, xp)
                                        if itens_check:
                                            try:
                                                el_check = itens_check[0].find_element(By.XPATH, "..")
                                            except:
                                                el_check = itens_check[0]
                                            if verificar_filtro_marcado(el_check):
                                                logger.info(f"    ✅ '{nome}' marcada (método 1)")
                                                return True

                                    # Método 2: JS click direto
                                    try:
                                        logger.debug(f"      [2/5] JS click direto...")
                                        self.driver.execute_script("arguments[0].click();", el)
                                        time.sleep(1.2)
                                        itens_check = self.driver.find_elements(By.XPATH, xp)
                                        if itens_check:
                                            try:
                                                el_check = itens_check[0].find_element(By.XPATH, "..")
                                            except:
                                                el_check = itens_check[0]
                                            if verificar_filtro_marcado(el_check):
                                                logger.info(f"    ✅ '{nome}' marcada (método 2)")
                                                return True
                                    except Exception as e:
                                        logger.debug(f"      Método 2 falhou: {e}")

                                    # Método 3: ActionChains
                                    try:
                                        logger.debug(f"      [3/5] ActionChains...")
                                        actions = ActionChains(self.driver)
                                        actions.move_to_element(el).pause(0.3).click().perform()
                                        time.sleep(1.2)
                                        itens_check = self.driver.find_elements(By.XPATH, xp)
                                        if itens_check:
                                            try:
                                                el_check = itens_check[0].find_element(By.XPATH, "..")
                                            except:
                                                el_check = itens_check[0]
                                            if verificar_filtro_marcado(el_check):
                                                logger.info(f"    ✅ '{nome}' marcada (método 3)")
                                                return True
                                    except Exception as e:
                                        logger.debug(f"      Método 3 falhou: {e}")

                                    # Método 4: Clique no span interno
                                    try:
                                        logger.debug(f"      [4/5] Span interno...")
                                        span_interno = el.find_element(By.XPATH, ".//span[contains(@class,'nome-item')]")
                                        self.driver.execute_script("arguments[0].click();", span_interno)
                                        time.sleep(1.2)
                                        itens_check = self.driver.find_elements(By.XPATH, xp)
                                        if itens_check:
                                            try:
                                                el_check = itens_check[0].find_element(By.XPATH, "..")
                                            except:
                                                el_check = itens_check[0]
                                            if verificar_filtro_marcado(el_check):
                                                logger.info(f"    ✅ '{nome}' marcada (método 4)")
                                                return True
                                    except Exception as e:
                                        logger.debug(f"      Método 4 falhou: {e}")

                                    # Método 5: Dispatch MouseEvent
                                    try:
                                        logger.debug(f"      [5/5] Dispatch MouseEvent...")
                                        self.driver.execute_script("""
                                            var el = arguments[0];
                                            var evt = new MouseEvent('click', {
                                                view: window, bubbles: true, cancelable: true
                                            });
                                            el.dispatchEvent(evt);
                                        """, el)
                                        time.sleep(1.2)
                                        itens_check = self.driver.find_elements(By.XPATH, xp)
                                        if itens_check:
                                            try:
                                                el_check = itens_check[0].find_element(By.XPATH, "..")
                                            except:
                                                el_check = itens_check[0]
                                            if verificar_filtro_marcado(el_check):
                                                logger.info(f"    ✅ '{nome}' marcada (método 5)")
                                                return True
                                    except Exception as e:
                                        logger.debug(f"      Método 5 falhou: {e}")

                                    logger.debug(f"      ❌ Todos os 5 métodos falharam para XPath {idx+1}")
                            except Exception as e:
                                logger.debug(f"    Erro tentativa {tentativa+1} XPath {idx+1}: {e}")
                                continue

                        # FALLBACK FINAL: Busca por texto exato via JavaScript (última tentativa)
                        if tentativa == 4:
                            try:
                                logger.info(f"    🔧 FALLBACK FINAL: Busca JavaScript por texto exato...")
                                for lab in labels:
                                    js_find_and_click = f"""
                                    var elementos = document.querySelectorAll('div.filtro-item, span.nome-item, div[class*="item"]');
                                    var encontrado = false;
                                    for (var i = 0; i < elementos.length; i++) {{
                                        var el = elementos[i];
                                        var texto = el.innerText || el.textContent;
                                        if (texto && texto.trim() === '{lab}') {{
                                            var clicavel = el.tagName === 'DIV' ? el : el.parentElement;
                                            if (clicavel) {{
                                                clicavel.click();
                                                encontrado = true;
                                                break;
                                            }}
                                        }}
                                    }}
                                    return encontrado;
                                    """
                                    resultado = self.driver.execute_script(js_find_and_click)
                                    if resultado:
                                        logger.info(f"    ⚡ JS encontrou e clicou em '{lab}'")
                                        time.sleep(1.5)
                                        # Valida
                                        if candidatos:
                                            itens_final = self.driver.find_elements(By.XPATH, candidatos[0])
                                            if itens_final:
                                                try:
                                                    el_final = itens_final[0].find_element(By.XPATH, "..")
                                                except:
                                                    el_final = itens_final[0]
                                                if verificar_filtro_marcado(el_final):
                                                    logger.info(f"    ✅ '{nome}' marcada (JS final)")
                                                    return True
                            except Exception as e:
                                logger.debug(f"    Erro fallback JS final: {e}")

                        if tentativa < 4:
                            time.sleep(0.7)  # Pausa maior entre tentativas

            logger.warning(f"  ✗ FALHA: Turma '{nome}' não pôde ser marcada após todas tentativas")
            return False

        # Caminho do DOCX de saída (mesmo diretório já usado pelo seu projeto)
        workspace_dir = os.path.dirname(os.path.abspath(__file__))
        default_docx = os.path.join(workspace_dir, "Diario_J_TST_com_variaveis.docx")
        docx_path = os.environ.get("JT_DOCX_PATH", default_docx)
        try:
            abs_docx = os.path.abspath(docx_path)
            logger.info(f"DOCX de saída configurado: {abs_docx}")
            pasta_doc = os.path.dirname(abs_docx) or workspace_dir
            if not os.path.isdir(pasta_doc):
                logger.warning(f"Diretório do DOCX não existe (vai ser criado): {pasta_doc}")
                os.makedirs(pasta_doc, exist_ok=True)
            else:
                logger.info(f"Diretório do DOCX OK: {pasta_doc}")
        except Exception:
            pass

        # ========== SELEÇÃO DE TRTs PELO USUÁRIO ==========
        logger.info("\n" + "="*80)
        logger.info("SELEÇÃO DE TRIBUNAIS")
        logger.info("="*80)

        # Verifica se há variável de ambiente para modo não-interativo
        modo_auto = os.environ.get("JT_AUTO_MODE", "0") == "1"

        if modo_auto:
            logger.info("Modo automático ativado (JT_AUTO_MODE=1)")
            trts_escolhidos = ["TRT3", "TRT24"]
            logger.info(f"Executando TRTs: {', '.join(trts_escolhidos)}")
        else:
            print("\n" + "="*80)
            print("  ESCOLHA OS TRIBUNAIS PARA PROCESSAR")
            print("="*80)
            print("\nOpções disponíveis:")
            print("  1 - TRT3 apenas (11 turmas)")
            print("  2 - TRT24 apenas (2 turmas)")
            print("  3 - Ambos (TRT3 + TRT24)")
            print("  0 - Sair/Cancelar")
            print("-"*80)

            while True:
                try:
                    escolha = input("\nDigite sua escolha (0-3): ").strip()

                    if escolha == "0":
                        logger.info("❌ Execução cancelada pelo usuário")
                        self.fail("Execução cancelada")
                    elif escolha == "1":
                        trts_escolhidos = ["TRT3"]
                        logger.info("✓ Selecionado: TRT3 apenas")
                        break
                    elif escolha == "2":
                        trts_escolhidos = ["TRT24"]
                        logger.info("✓ Selecionado: TRT24 apenas")
                        break
                    elif escolha == "3":
                        trts_escolhidos = ["TRT3", "TRT24"]
                        logger.info("✓ Selecionado: Ambos (TRT3 e TRT24)")
                        break
                    else:
                        print("❌ Opção inválida! Digite 0, 1, 2 ou 3.")
                except KeyboardInterrupt:
                    print("\n\n❌ Execução interrompida pelo usuário")
                    logger.info("Execução interrompida (Ctrl+C)")
                    self.fail("Execução interrompida")
                except Exception as e:
                    print(f"❌ Erro: {e}")
                    print("Tente novamente.")

        print("\n" + "="*80)
        print(f"  PROCESSANDO: {', '.join(trts_escolhidos)}")
        print("="*80 + "\n")

        # NOTA: turmas será definido dinamicamente por TRT (11 para TRT3, 2 para TRT24)
        trts = trts_escolhidos  # Usa a escolha do usuário
        turmas = []  # Será preenchido dinamicamente para cada TRT
        total_geral = 0

        # Função que executa a extração usando a lógica atual, com alvo local
        def extrair_para_turma(alvo_local: int = 10) -> int:
            try:
                # Aguarda resultados carregarem
                if not self._wait_results_loaded(30):
                    logger.warning("Prosseguindo com o que estiver visível.")
                # Revelar cartões
                try:
                    _ = self._reveal_all_cards_on_page(max_scrolls=15, settle_cycles=2)
                except Exception:
                    pass
                inseridos = 0
                vistos_ids = set()
                numeros_processados = set()
                tentativas = 0
                max_tentativas = 90
                # Controle de páginas
                pagina_atual = 1
                max_paginas = 10
                while tentativas < max_tentativas and pagina_atual <= max_paginas:
                    cards = self._buscar_cartoes()
                    if not cards:
                        self._scroll_to_bottom(tries=2)
                        time.sleep(0.5)
                        cards = self._buscar_cartoes()
                        if not cards:
                            logger.info("Nenhum cartão nesta página.")
                    for proximo in cards:
                        try:
                            try:
                                cid = proximo.id
                            except Exception:
                                cid = None
                            if cid and cid in vistos_ids:
                                continue
                            self._scroll_center(proximo)
                            self._hover(proximo)
                            time.sleep(0.1)
                            header_lines = self._extrair_cabecalho(proximo)
                            if not header_lines:
                                self._scroll_by(0.3)
                                time.sleep(0.2)
                                header_lines = self._extrair_cabecalho(proximo)
                            if not header_lines:
                                if cid:
                                    vistos_ids.add(cid)
                                tentativas += 1
                                continue
                            # Filtrar seções de dissídios
                            try:
                                banned_terms = [
                                    "1ª seção de dissídios individuais",
                                    "2ª seção de dissídios individuais",
                                    "seção de dissídios coletivos",
                                ]
                                def _norm_txt(s):
                                    return re.sub(r"\s+", " ", (s or "")).strip().lower()
                                hl_norm = [_norm_txt(x) for x in header_lines]
                                if any(any(bt in h for bt in banned_terms) for h in hl_norm):
                                    if cid:
                                        vistos_ids.add(cid)
                                    tentativas += 1
                                    continue
                            except Exception:
                                pass
                            dados, tribunal_tag = self._montar_dados_a_partir_do_cabecalho(header_lines)
                            num = dados.get('numero_processo') or ''
                            if num and num in numeros_processados:
                                if cid:
                                    vistos_ids.add(cid)
                                tentativas += 1
                                continue
                            ementa = self._obter_ementa(proximo)
                            # Usa caminho real persistido se já houve fallback por arquivo bloqueado
                            destino_docx = getattr(self, '_docx_real_path', docx_path)
                            self._append_to_docx(destino_docx, dados, header_lines, ementa)
                            inseridos += 1
                            if cid:
                                vistos_ids.add(cid)
                            if num:
                                numeros_processados.add(num)
                            self._scroll_by(0.45)
                            time.sleep(0.25)
                        except Exception as e:
                            logger.error(f"Erro processando cartão {inseridos+1}: {e}")
                            logger.debug(traceback.format_exc())
                            try:
                                cid = proximo.id
                                vistos_ids.add(cid)
                            except Exception:
                                pass
                            tentativas += 1
                    # Sempre tentar avançar para a próxima página até o limite
                    self._reveal_all_cards_on_page(max_scrolls=2, settle_cycles=1)
                    self._scroll_to_bottom(tries=2)
                    # Checa limite de páginas
                    if pagina_atual >= max_paginas:
                        logger.info(f"Limite de páginas atingido ({max_paginas}). Interrompendo paginação da turma.")
                        break
                    if not self._go_to_next_page():
                        logger.info("Não foi possível avançar para a próxima página. Finalizando paginação.")
                        break
                    pagina_atual += 1
                    logger.info(f"📄 Avançou para a página {pagina_atual}")
                    time.sleep(0.5)
                logger.info(f"Blocos inseridos para a turma: {inseridos}")
                return inseridos
            except Exception:
                return 0

        # Loop TRT -> Turmas sem recarregar página: alterna filtros no mesmo carregamento
        # Abre a página uma única vez e aplica filtros básicos no início
        self.open(url)
        self.wait_for_element("body", timeout=30)
        time.sleep(1.0)
        aplicar_filtros_basicos()

        # Configura paginação para 10 itens UMA ÚNICA VEZ (não precisa reajustar para cada turma)
        logger.info("⚙️ Configurando paginação para 10 itens por página...")
        try:
            self._ajustar_itens_por_pagina('10')
            logger.info("✓ Paginação configurada")
        except Exception as e:
            logger.warning(f"⚠ Falha ao configurar paginação: {e}")

        trt_atual = None

        # Abrir instância persistente do Word COM para atualizações de sumário por turma
        if not getattr(self, 'skip_sumario', False):
            opened = self._open_word_app()
            if opened:
                logger.info("Word COM aberto para atualizações persistentes de sumário.")
            else:
                logger.warning("Não foi possível abrir Word COM; atualizações persistentes de sumário serão ignoradas.")

        def desmarcar_trt_atual():
            """Desmarca o TRT atualmente selecionado clicando no X do chip"""
            try:
                logger.info("  🧹 Desmarcando TRT anterior...")
                # XPath para encontrar chip de TRT (não turma)
                xp_remove_trt = (
                    "//div[contains(@class,'p-chip')]"
                    "[.//div[contains(@class,'p-chip-text') and contains(text(),'TRT')]]"
                    "//span[contains(@class,'pi-chip-remove-icon') and contains(@class,'pi-times-circle')]"
                )

                removes = self.driver.find_elements(By.XPATH, xp_remove_trt)
                removidos = 0

                for rem in removes:
                    try:
                        if rem.is_displayed():
                            self._scroll_center(rem)
                            time.sleep(0.2)
                            self._safe_js_click(rem)
                            removidos += 1
                            time.sleep(0.6)
                            logger.info(f"    ✓ TRT desmarcado")
                    except Exception:
                        continue

                return removidos > 0
            except Exception as e:
                logger.debug(f"    Erro ao desmarcar TRT: {e}")
                return False

        for trt in trts:
            # Desmarca TRT anterior antes de selecionar novo
            if trt_atual is not None:
                desmarcar_trt_atual()
                time.sleep(0.5)

            logger.info(f"🏛️ Alternando para {trt}...")
            selecionar_trt(trt)

            # Espera mais tempo para garantir que o TRT está marcado e estável
            time.sleep(1.2)
            self._wait_results_loaded(30)

            # Pequeno scroll após selecionar TRT para posicionar a lista
            try:
                self._scroll_by(0.2)
                time.sleep(0.3)
            except Exception:
                pass

            # Tentar revelar cartões para garantir que resultados apareçam
            try:
                self._reveal_all_cards_on_page(max_scrolls=3, settle_cycles=1)
            except Exception:
                pass

            trt_atual = trt

            # Define número de turmas baseado no TRT
            if trt == "TRT24":
                turmas = [f"{i:02d}ª Turma" for i in range(1, 3)]  # 2 turmas (1ª e 2ª)
                logger.info(f"🏛️ {trt} possui 2 turmas")
            else:  # TRT3 ou outros
                turmas = [f"{i:02d}ª Turma" for i in range(1, 12)]  # 11 turmas
                logger.info(f"🏛️ {trt} possui 11 turmas")

            # Itera turmas alternando a seleção sem reload
            for turma in turmas:
                # Marca a nova turma
                logger.info(f"🔄 Selecionando {turma} no {trt}...")
                selecionar_turma(turma)
                time.sleep(0.8)
                self._wait_results_loaded(25)

                # Pequeno scroll após selecionar a Turma antes de extrair
                try:
                    self._scroll_by(0.25)
                    time.sleep(0.25)
                except Exception:
                    pass

                # Tentar revelar cartões já aqui
                try:
                    self._reveal_all_cards_on_page(max_scrolls=3, settle_cycles=1)
                except Exception:
                    pass

                # Extrair dados
                logger.info(f"📄 Extraindo primeira página de {turma} ({trt})...")
                qtd_extraida = extrair_para_turma(alvo_local=10)
                total_geral += qtd_extraida
                logger.info(f"✓ Concluída extração de {qtd_extraida} processos para {turma}.")

                # Atualizar sumário com páginas reais ao fim de cada turma
                if qtd_extraida > 0 and not getattr(self, 'skip_sumario', False):
                    try:
                        logger.info(f"📖 Atualizando paginação real no sumário para {turma}...")
                        # 1) Garante entradas via PAGEREF
                        try:
                            self._atualizar_sumario_com_pageref(docx_path, self._turma_bookmarks)
                        except Exception as e1:
                            logger.debug(f"Falha ao inserir PAGEREF: {e1}")
                        # 2) Atualiza campos com Word COM (instância persistente)
                        if not self._atualizar_sumario_win32(docx_path):
                            logger.warning("⚠ Não foi possível calcular páginas via Word COM nesta etapa.")
                        else:
                            logger.info("✅ Paginação atualizada com sucesso no arquivo.")
                    except Exception as e_sum:
                        logger.error(f"Erro ao atualizar sumário parcial: {e_sum}")

        logger.info(f"Concluído. Blocos inseridos no DOCX: {total_geral}")
        assert total_geral > 0, "Nenhum bloco foi inserido no documento."

        # ATUALIZAÇÃO FINAL DO SUMÁRIO E LIMPEZA DO WORD
        try:
            if not getattr(self, 'skip_sumario', False):
                self._atualizar_sumario_win32(docx_path)
        except Exception as e:
            logger.warning(f"Atualização final do sumário falhou: {e}")

        # LIMPEZA FINAL: FECHAR O WORD (se aberto por _get_word_app)
        try:
            if getattr(self, '_word_app', None):
                try:
                    try:
                        self._word_app.ScreenUpdating = True
                    except Exception:
                        pass
                    self._word_app.Quit()
                except Exception:
                    pass
                self._word_app = None
                self._word_persistent = False
                logger.info("Aplicação Word encerrada com sucesso.")
        except Exception:
            pass

        # Opcional: abrir pasta de saída automaticamente (defina JT_OPEN_FOLDER=1)
        try:
            if os.environ.get("JT_OPEN_FOLDER", "0") == "1":
                pasta = os.path.dirname(docx_path)
                if os.path.isdir(pasta):
                    logger.info(f"Abrindo pasta de saída: {pasta}")
                    try:
                        os.startfile(pasta)  # Windows
                    except Exception:
                        pass
        except Exception:
            pass

    # ---------- Captura dos Cartões ----------
    def _buscar_cartoes_beautifulsoup(self) -> List:
        """Método alternativo usando BeautifulSoup para extrair cartões."""
        if BeautifulSoup is None:
            return []
        
        try:
            html = self.driver.page_source
            soup = BeautifulSoup(html, 'html.parser')
            
            # Procurar por divs que contenham botões de copiar ementa
            cartoes_encontrados = []
            
            # Estratégia 1: Procurar por elementos com classe 'doc-card'
            cards = soup.find_all('div', class_=lambda x: x and 'doc-card' in x)
            logger.info(f"BeautifulSoup: Encontrados {len(cards)} elementos com 'doc-card'")
            
            if not cards:
                # Estratégia 2: Procurar por elementos que contenham o botão de copiar
                buttons = soup.find_all('button', attrs={'aria-label': 'Botão copiar ementa do acórdão'})
                logger.info(f"BeautifulSoup: Encontrados {len(buttons)} botões de copiar ementa")
                
                # Para cada botão, subir na árvore para encontrar o cartão pai
                for btn in buttons:
                    # Subir até encontrar um container apropriado
                    parent = btn.parent
                    nivel = 0
                    while parent and nivel < 10:
                        # Verificar se tem classe que indique ser um cartão/documento
                        classes = parent.get('class', [])
                        classes_str = ' '.join(classes) if isinstance(classes, list) else str(classes)
                        
                        if any(keyword in classes_str.lower() for keyword in ['doc-card', 'documento', 'card', 'result']):
                            # Converter de volta para elemento Selenium
                            try:
                                # Usar um atributo único para encontrar o elemento
                                if parent.get('id'):
                                    elem = self.driver.find_element(By.ID, parent.get('id'))
                                    cartoes_encontrados.append(elem)
                                    break
                                elif classes:
                                    # Tentar encontrar por classe
                                    xpath = f"//{parent.name}[contains(@class, '{classes[0]}')]"
                                    elems = self.driver.find_elements(By.XPATH, xpath)
                                    if elems:
                                        cartoes_encontrados.append(elems[0])
                                        break
                            except Exception:
                                pass
                        
                        parent = parent.parent
                        nivel += 1
            else:
                # Converter elementos BeautifulSoup para Selenium
                for card in cards:
                    try:
                        if card.get('id'):
                            elem = self.driver.find_element(By.ID, card.get('id'))
                            cartoes_encontrados.append(elem)
                        else:
                            classes = card.get('class', [])
                            if classes:
                                xpath = f"//div[contains(@class, '{classes[0]}')]"
                                elems = self.driver.find_elements(By.XPATH, xpath)
                                if elems:
                                    cartoes_encontrados.append(elems[0])
                    except Exception:
                        continue
            
            logger.info(f"BeautifulSoup: Retornando {len(cartoes_encontrados)} cartões")
            return cartoes_encontrados
            
        except Exception as e:
            logger.error(f"Erro ao usar BeautifulSoup: {e}")
            return []
    
    def _buscar_cartoes(self) -> List:
        sel = self.selectors.get('jt', {}).get('resultados', {})
        caminhos = sel.get('cartao_root', [])
        
        # Debug: Verificar se há elementos na página
        if os.environ.get("JT_DEBUG_LOG", "0") == "1":
            try:
                # Tentar encontrar qualquer div com 'documento' ou 'card' no class
                debug_divs = self.driver.find_elements(By.XPATH, "//div[contains(@class, 'documento') or contains(@class, 'card') or contains(@class, 'result')]")
                logger.info(f"DEBUG: Encontrados {len(debug_divs)} divs com 'documento/card/result' no class")
                
                # Verificar estrutura geral da página
                all_divs = self.driver.find_elements(By.XPATH, "//div[@class]")
                logger.info(f"DEBUG: Total de divs com classe na página: {len(all_divs)}")
                
                # Mostrar algumas classes para entender a estrutura
                if all_divs:
                    classes_sample = set()
                    for d in all_divs[:50]:  # Primeiros 50
                        cls = d.get_attribute('class')
                        if cls and ('doc' in cls.lower() or 'card' in cls.lower() or 'result' in cls.lower()):
                            classes_sample.add(cls)
                    if classes_sample:
                        logger.info(f"DEBUG: Classes relevantes encontradas: {list(classes_sample)[:10]}")
            except Exception as e:
                logger.error(f"DEBUG: Erro ao verificar elementos: {e}")
        
        vistos = []
        vistos_ids = set()
        # 1) Tentar encontrar contêineres completos de cartões
        for xp in caminhos:
            try:
                elems = self.driver.find_elements(By.XPATH, xp)
                for e in elems:
                    try:
                        if not e.is_displayed():
                            continue
                    except Exception:
                        continue
                    key = e.id
                    if key not in vistos_ids:
                        vistos.append(e)
                        vistos_ids.add(key)
            except Exception:
                continue
        if vistos:
            return vistos
        
        # Debug: Salvar HTML se não encontrou nada
        if os.environ.get("JT_DEBUG_LOG", "0") == "1":
            try:
                html_content = self.driver.page_source
                debug_file = os.path.join(os.path.dirname(__file__), "debug_jt_no_cards.html")
                with open(debug_file, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                logger.info(f"DEBUG: HTML da página salvo em {debug_file}")
            except Exception as e:
                logger.error(f"DEBUG: Erro ao salvar HTML: {e}")
        
        # 2) Fallback: localizar diretamente as seções de cabeçalho e subir ao ancestral com botão copiar
        header_fb_paths = sel.get('header_section_fallback', [])
        candidatos = []
        for xp in header_fb_paths:
            try:
                secs = self.driver.find_elements(By.XPATH, xp)
                for sec in secs:
                    try:
                        if not sec.is_displayed():
                            continue
                    except Exception:
                        continue
                    candidatos.append(sec)
            except Exception:
                continue
        for sec in candidatos:
            try:
                # tentar pegar ancestral que contenha o botão de copiar
                ancestors = sec.find_elements(By.XPATH, "ancestor::*[.//span[contains(@class,'doc-botao-icone') and normalize-space()='content_copy']]")
                card = ancestors[0] if ancestors else sec
                key = card.id
                if key not in vistos_ids:
                    vistos.append(card)
                    vistos_ids.add(key)
            except Exception:
                continue
        logger.info(f"Cartões via fallback de seção: {len(vistos)}")
        
        # 3) Fallback final: usar BeautifulSoup
        if not vistos and BeautifulSoup is not None:
            logger.info("Tentando extrair cartões com BeautifulSoup...")
            vistos = self._buscar_cartoes_beautifulsoup()
        
        return vistos

    # ---------- Extração do Cabeçalho ----------
    def _extrair_cabecalho(self, card) -> List[str]:
        sel = self.selectors.get('jt', {}).get('resultados', {})
        cab_paths = sel.get('cabecalho_section', [])
        for rel in cab_paths:
            try:
                secs = card.find_elements(By.XPATH, rel)
                for sec in secs:
                    try:
                        if not sec.is_displayed():
                            continue
                    except Exception:
                        continue
                    linhas = sec.find_elements(By.XPATH, sel.get('cabecalho_linhas', ".//div[contains(@class,'doc-texto') ]"))
                    textos = []
                    for l in linhas:
                        try:
                            t = (l.text or '').strip()
                            if t:
                                textos.append(re.sub(r"\s+", " ", t))
                        except Exception:
                            continue
                    if textos:
                        return textos
            except Exception:
                continue
        # Fallback: se o próprio 'card' for uma seção com as linhas
        try:
            linhas = card.find_elements(By.XPATH, ".//div[contains(@class,'doc-texto')]")
            textos = []
            for l in linhas:
                try:
                    t = (l.text or '').strip()
                    if t:
                        textos.append(re.sub(r"\s+", " ", t))
                except Exception:
                    continue
            if textos:
                return textos
        except Exception:
            pass
        return []

    def _montar_dados_a_partir_do_cabecalho(self, linhas: List[str]) -> Tuple[dict, Optional[str]]:
        numero = ''
        orgao = ''
        relator = ''
        publicacao = ''
        tipo_doc = ''
        tribunal_tag = None

        # 1) número CNJ na primeira/segunda linha
        cnj_re = re.compile(r"\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}")
        for base in linhas[:2]:
            m = cnj_re.search(base)
            if m:
                numero = m.group(0)
                break

        # 2) tribunal (TRTn) e tipo de doc (segunda linha costuma ser o tipo)
        if linhas:
            # Primeira linha contém ex.: "TRT3 - ROT <número>"
            trib_m = re.search(r"\bTRT\d{1,2}\b", linhas[0])
            if trib_m:
                tribunal_tag = trib_m.group(0)
        if len(linhas) >= 2:
            tipo_doc = linhas[1].strip()

        # 3) relatoria
        for s in linhas:
            n = s.strip()
            if re.search(r"\bRelatoria de\b", n, re.IGNORECASE):
                relator = re.sub(r".*Relatoria de\s*", "", n, flags=re.IGNORECASE).strip()
                break

        # 4) turma (para Órgão Judicante)
        for s in linhas:
            n = s.strip()
            if re.search(r"\b\d{1,2}ª\s*Turma\b", n, re.IGNORECASE):
                orgao = n
                break
        if not orgao and tribunal_tag:
            orgao = tribunal_tag

        # 5) data de juntada -> mapeamos como Publicação
        for s in linhas:
            n = s.strip()
            if re.search(r"Juntado aos autos em", n, re.IGNORECASE):
                dt = re.sub(r".*Juntado aos autos em\s*", "", n, flags=re.IGNORECASE).strip()
                publicacao = dt
                break

        dados = {
            'numero_processo': numero,
            'referencias': {
                'Órgão Judicante': orgao,
                'Relator': relator,
                'Julgamento': '',
                'Publicação': publicacao,
                'Tipo de Documento': tipo_doc,
            }
        }
        return dados, tribunal_tag

    # ---------- Ementa ----------
    def _obter_ementa(self, card) -> str:
        """Obtém ementa usando apenas o botão de copiar (clipboard)."""
        if self.disable_clipboard:
            logger.warning("Clipboard desabilitado por variável de ambiente.")
            return ''
        
        if pyperclip is None:
            logger.warning("pyperclip não disponível")
            return ''
        
        # Encontrar botão de copiar
        try:
            self._scroll_center(card)
            self._hover(card)
            time.sleep(0.05)
        except Exception:
            pass
        
        el_copy = self._find_copy_element(card)
        if not el_copy:
            self._scroll_by(0.25)
            time.sleep(0.2)
            el_copy = self._find_copy_element(card)
        
        if not el_copy:
            logger.warning("Botão de copiar ementa não encontrado para este cartão.")
            return ''
        
        # Preparar e clicar no botão
        try:
            self._scroll_center(el_copy)
            self._hover(el_copy)
            time.sleep(0.1)
            
            # Limpar clipboard antes
            try:
                pyperclip.copy('')
            except Exception:
                pass
            
            # Clicar no botão de copiar
            if not self._safe_js_click(el_copy):
                logger.warning("Clique no botão de copiar falhou.")
                return ''
            
            # Aguardar mais tempo para garantir que o texto foi copiado
            time.sleep(2.0)
            
            # Tentar múltiplas vezes se necessário
            max_tentativas = 3
            ementa = ''
            
            for tentativa in range(max_tentativas):
                try:
                    ementa = pyperclip.paste() or ''
                    if ementa and ementa.strip() and len(ementa.strip()) > 50:
                        break
                    if tentativa < max_tentativas - 1:
                        logger.warning(f"Tentativa {tentativa + 1}: clipboard vazio ou muito curto, aguardando...")
                        time.sleep(1.5)
                except Exception as e:
                    logger.error(f"Erro ao obter texto do clipboard (tentativa {tentativa + 1}): {e}")
                    if tentativa < max_tentativas - 1:
                        time.sleep(1.5)
            
            if ementa and ementa.strip():
                # Normalizar quebras de linha
                ementa = ementa.replace('\r\n', '\n').replace('\r', '\n')
                ementa = ementa.strip()

                # LIMPEZA DE HTML: Remove tags HTML se presentes
                if '<div' in ementa or '<section' in ementa or '<button' in ementa:
                    logger.warning("⚠ Detectado HTML no clipboard. Limpando...")

                    # Usa BeautifulSoup se disponível
                    if BeautifulSoup:
                        try:
                            soup = BeautifulSoup(ementa, 'html.parser')

                            # Remove elementos indesejados
                            for tag in soup.find_all(['button', 'script', 'style']):
                                tag.decompose()

                            # Pega apenas o texto
                            ementa = soup.get_text(separator=' ', strip=True)
                            logger.info(f"✓ HTML limpo com BeautifulSoup")
                        except Exception as e:
                            logger.warning(f"Erro ao limpar HTML com BeautifulSoup: {e}")
                            # Fallback: regex simples
                            ementa = re.sub(r'<[^>]+>', '', ementa)
                            ementa = re.sub(r'\s+', ' ', ementa).strip()
                    else:
                        # Fallback: regex simples para remover tags
                        ementa = re.sub(r'<[^>]+>', '', ementa)
                        ementa = re.sub(r'\s+', ' ', ementa).strip()

                # Remove o cabeçalho "Acórdão" se presente
                ementa = re.sub(r'^Acórdão\s*', '', ementa, flags=re.IGNORECASE)

                # Remove "Inteiro teor" e variações (incluindo variações com parênteses e sem)
                ementa = re.sub(r'Inteiro teor\s*\([^\)]*\)', '', ementa, flags=re.IGNORECASE)
                ementa = re.sub(r'Inteiro teor[^\n\.]*', '', ementa, flags=re.IGNORECASE)
                ementa = re.sub(r'\s+Inteiro\s+teor\s*', ' ', ementa, flags=re.IGNORECASE)

                # Remove botão "ler inteiro teor" se presente
                ementa = re.sub(r'\s*ler inteiro teor\s*,?\s*', ' ', ementa, flags=re.IGNORECASE)

                # Remove fragmentos de texto cortado que aparecem antes da ementa (ex: "...1731-25.2010.5.24.0022")
                ementa = re.sub(r'\.{3,}\d{4}-\d{2}\.\d{4}\.\d+\.\d{2}\.\d{4}[^\n]*', '', ementa)

                # Procurar pela seção "Ementa:" e pegar apenas o que vem depois
                # O clipboard traz: cabeçalho + possível texto do acórdão + "Ementa: \n" + texto da ementa
                match = re.search(r'Ementa:\s*\n(.+)', ementa, re.DOTALL | re.IGNORECASE)
                if match:
                    ementa = match.group(1).strip()
                    logger.info(f"✓ Ementa extraída (após 'Ementa:'): {len(ementa)} caracteres")
                else:
                    # Tenta encontrar "Ementa:" sem quebra de linha
                    match2 = re.search(r'Ementa:\s*(.+)', ementa, re.DOTALL | re.IGNORECASE)
                    if match2:
                        ementa = match2.group(1).strip()
                        logger.info(f"✓ Ementa extraída (após 'Ementa:' inline): {len(ementa)} caracteres")
                    else:
                        logger.info(f"✓ Ementa obtida via clipboard (sem marcador 'Ementa:'): {len(ementa)} caracteres")

                # Limpa espaços múltiplos e reticências cortadas
                ementa = re.sub(r'\s+', ' ', ementa).strip()
                ementa = re.sub(r'\.{4,}', '...', ementa)  # Normaliza reticências

                # Remove texto cortado no início/fim (ex: "...uando")
                ementa = re.sub(r'^\.{3,}[a-z]+\s+', '', ementa)  # Remove início cortado
                ementa = re.sub(r'\s+[a-z]+\.{3,}$', '', ementa)  # Remove fim cortado

                # NÃO remove informações do tribunal - mantém tudo da ementa
                # Remove apenas limpeza final de espaços múltiplos
                ementa = re.sub(r'\s+', ' ', ementa).strip()

                # Debug: mostrar início e fim da ementa
                if os.environ.get("JT_DEBUG_LOG", "0") == "1":
                    logger.info(f"Início da ementa limpa: {ementa[:150]}...")
                    logger.info(f"Fim da ementa limpa: ...{ementa[-150:]}")

                return ementa
            else:
                logger.warning("Clipboard vazio ou muito curto após todas as tentativas.")
                return ''
                
        except Exception as e:
            logger.error(f"Erro ao copiar ementa: {e}")
            return ''

    # ---------- DOCX (mantendo formatação) ----------
    def _append_to_docx(self, doc_path: str, dados: dict, header_lines: List[str], ementa: str):
        try:
            # ---- DEBUG LOGGING ----
            if os.environ.get("JT_DEBUG_LOG", "0") == "1":
                logger.info("---- DADOS PARA GRAVAÇÃO ----")
                logger.info(f"HEADER: {header_lines}")
                logger.info(f"DADOS: {dados}")
                logger.info(f"EMENTA (primeiros 300 chars): {ementa[:300] if ementa else '<<VAZIA>>'}")
                logger.info("--------------------------")
            # ---- FIM DEBUG LOGGING ----

            logger.info(f"Gravando bloco no DOCX: {os.path.abspath(doc_path)}")
            # Garante documento com 'Sumário' no topo quando for novo
            if os.path.exists(doc_path):
                doc = Document(doc_path)
            else:
                doc = self._prepare_document_with_sumario(doc_path)
            # Garantir que haja 'Sumário' mesmo em documento existente
            _ = self._prepare_document_with_sumario(doc_path)
            doc = Document(doc_path)

            # Localizar parágrafo 'Sumário' para inserir ANTES dele
            suminfo = self._buscar_sumario_em_documento(doc)
            anchor_para = suminfo['elemento'] if suminfo else None
            if anchor_para is None:
                # Fallback: se não encontrou, cria e recarrega
                doc = self._prepare_document_with_sumario(doc_path)
                doc = Document(doc_path)
                suminfo = self._buscar_sumario_em_documento(doc)
                anchor_para = suminfo['elemento'] if suminfo else None

            # Se for a primeira ocorrência de uma Turma/Bloco, criar heading e bookmark ANTES do Sumário
            try:
                orgao = (dados or {}).get('referencias', {}).get('Órgão Judicante') or (dados or {}).get('referencias', {}).get('Orgão Judicante') or ''
                ident = self._extrair_id_bloco(orgao)
                logger.info(f"🔍 Órgão Judicante: {orgao} → Identificador: {ident}")
            except Exception as e:
                logger.warning(f"Erro ao extrair ID do bloco: {e}")
                ident = None
            if ident and ident not in self._turma_bookmarks and anchor_para is not None:
                titulo = self._descricao_por_identificador(ident)
                p_head = anchor_para.insert_paragraph_before(titulo)
                try:
                    p_head.style = 'Heading 1'
                except Exception:
                    pass
                bm_name = self._sanitizar_nome_bookmark(f"BM_TURMA_{ident}")
                self._inserir_bookmark_no_paragrafo(doc, p_head, bm_name)
                self._turma_bookmarks[ident] = bm_name
                logger.info(f"📍 Bookmark criado: {ident} → {bm_name}")
                # separação após heading
                anchor_para.insert_paragraph_before('')

            # Inserir bloco (cabeçalho + 'Ementa:' + ementa) ANTES do Sumário
            if anchor_para is not None:
                # separador em branco antes do bloco
                anchor_para.insert_paragraph_before('')
                # Cabeçalho + rótulo 'Ementa:'
                linhas = list(header_lines or [])
                # Ementa
                linhas.append('Ementa:')  # Ementa
                for hl in linhas:
                    if hl is None:
                        continue
                    p = anchor_para.insert_paragraph_before('')
                    self._format_line(p, hl)
                # Ementa com hyperlinks inline
                p_em = anchor_para.insert_paragraph_before('')
                # Reusar lógica de hyperlink: construir diretamente no parágrafo
                try:
                    # Texto da ementa com espaço final para facilitar reconhecimento de URLs no Word
                    ementa_txt = (ementa or '').rstrip()
                    if ementa_txt:
                        ementa_txt = ementa_txt + ' '
                    # Implementação local: mesma usada em _add_ementa_with_inline_links
                    def add_text_run_local(pp, t: str):
                        if not t:
                            return
                        r = pp.add_run(t)
                        r.font.name = 'Arial MT'
                        r.font.size = Pt(8)
                    def add_hyperlink_run_local(pp, display: str, url: str):
                        r_id = pp.part.relate_to(url, RT.HYPERLINK, is_external=True)
                        hl = OxmlElement('w:hyperlink')
                        hl.set(qn('r:id'), r_id)
                        run = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')
                        rStyle = OxmlElement('w:rStyle')
                        rStyle.set(qn('w:val'), 'Hyperlink')
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), 'Arial MT')
                        rFonts.set(qn('w:hAnsi'), 'Arial MT')
                        rFonts.set(qn('w:cs'), 'Arial MT')
                        sz = OxmlElement('w:sz')
                        sz.set(qn('w:val'), '16')
                        rPr.append(rStyle)
                        rPr.append(rFonts)
                        rPr.append(sz)
                        run.append(rPr)
                        t = OxmlElement('w:t')
                        t.text = display
                        run.append(t)
                        hl.append(run)
                        pp._p.append(hl)
                    # Parágrafo formatado
                    pf = p_em.paragraph_format
                    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    pattern = re.compile(r'<(https?://[^>\s]+)>|(https?://\S+)')
                    pos = 0
                    texto = ementa_txt
                    for m in pattern.finditer(texto):
                        start, end = m.start(), m.end()
                        add_text_run_local(p_em, texto[pos:start])
                        url = m.group(1) or m.group(2)
                        bracketed = m.group(1) is not None
                        if bracketed:
                            add_text_run_local(p_em, '<')
                        add_hyperlink_run_local(p_em, url, url)
                        if bracketed:
                            add_text_run_local(p_em, '>')
                        pos = end
                    add_text_run_local(p_em, texto[pos:])
                except Exception:
                    # Fallback simples
                    self._format_line(p_em, ementa or '')
                # separador final
                anchor_para.insert_paragraph_before('')
            else:
                # Fallback: inserir no final como antes (caso raro)
                doc.add_paragraph('')
                linhas = list(header_lines or [])
                linhas.append('Ementa:')  # Ementa
                for hl in linhas:
                    if hl is None:
                        continue
                    p = doc.add_paragraph()
                    self._format_line(p, hl)
                ementa_txt = (ementa or '').rstrip()
                if ementa_txt:
                    ementa_txt = ementa_txt + ' '
                self._add_ementa_with_inline_links(doc, ementa_txt)
                doc.add_paragraph('')

            try:
                doc.save(doc_path)
            except PermissionError as pe:
                try:
                    base, ext = os.path.splitext(doc_path)
                    alt_path = f"{base}_{int(time.time())}{ext or '.docx'}"
                    logger.warning(f"Permissão negada ao salvar em {os.path.abspath(doc_path)} (arquivo pode estar aberto). Salvando cópia alternativa em: {os.path.abspath(alt_path)}")
                    doc.save(alt_path)
                    # Persistir caminho alternativo para usos futuros nesta sessão
                    self._docx_real_path = alt_path
                    doc_path = alt_path
                except Exception:
                    raise pe
            try:
                sz = os.path.getsize(doc_path)
                mt = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(os.path.getmtime(doc_path)))
                logger.info(f"DOCX salvo: tamanho={sz} bytes | mtime={mt}")
            except Exception:
                pass
            # Tentar atualizar o Sumário/TOC via Word COM imediatamente após salvar (por turma)
            try:
                try:
                    updated = self._atualizar_sumario_win32(doc_path)
                except Exception:
                    updated = False
                if updated:
                    logger.info("Sumário atualizado via Word COM após salvar o bloco.")
                else:
                    logger.debug("Atualização do Sumário via Word COM retornou False ou não foi possível.")
            except Exception:
                logger.debug("Falha silenciosa ao tentar atualizar sumário via Word COM.")
        except Exception as e:
            logger.error(f"Erro ao escrever no DOCX: {e}")
            logger.debug(traceback.format_exc())

    def _add_ementa_with_inline_links(self, doc: Document, texto: str):
        try:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE

            def add_text_run(t: str):
                if not t:
                    return
                r = p.add_run(t)
                r.font.name = 'Arial MT'
                r.font.size = Pt(8)

            def add_hyperlink_run(display: str, url: str):
                # cria relação externa
                r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
                hl = OxmlElement('w:hyperlink')
                hl.set(qn('r:id'), r_id)
                run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                # estilo Hyperlink (azul sublinhado)
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), 'Hyperlink')
                # Fonte e tamanho
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Arial MT')
                rFonts.set(qn('w:hAnsi'), 'Arial MT')
                rFonts.set(qn('w:cs'), 'Arial MT')
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '16')
                rPr.append(rStyle)
                rPr.append(rFonts)
                rPr.append(sz)
                run.append(rPr)
                t = OxmlElement('w:t')
                t.text = display
                run.append(t)
                hl.append(run)
                p._p.append(hl)

            # Regex: <url> ou url simples
            pattern = re.compile(r'<(https?://[^>\s]+)>|(https?://\S+)')
            pos = 0
            for m in pattern.finditer(texto):
                start, end = m.start(), m.end()
                add_text_run(texto[pos:start])
                url = m.group(1) or m.group(2)
                bracketed = m.group(1) is not None
                if bracketed:
                    add_text_run('<')
                add_hyperlink_run(url, url)
                if bracketed:
                    add_text_run('>')
                pos = end
            add_text_run(texto[pos:])
        except Exception:
            # fallback simples
            p = doc.add_paragraph(texto)
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = p.runs[0]
            run.font.name = 'Arial MT'
            run.font.size = Pt(8)
            return

    def _atualizar_sumario_win32(self, doc_path: str) -> bool:
        """
        Abre o DOCX, força a repaginação e atualiza campos (PAGEREF) no corpo.
        Mantém a instância do Word aberta para performance.
        """
        try:
            import win32com.client as win32

            abs_path = os.path.abspath(doc_path)
            if not os.path.exists(abs_path):
                return False

            # Reutiliza a instância rápida; não cria nova se indisponível
            word = self._get_word_app()
            if not word:
                return False

            try:
                doc_com = word.Documents.Open(
                    abs_path,
                    ReadOnly=False,
                    AddToRecentFiles=False,
                    Visible=False,
                )
            except Exception:
                time.sleep(0.6)
                doc_com = word.Documents.Open(
                    abs_path,
                    ReadOnly=False,
                    AddToRecentFiles=False,
                    Visible=False,
                )

            try:
                # Repagina para calcular páginas corretas
                try:
                    doc_com.Repaginate()
                except Exception:
                    pass

                # Atualiza apenas os campos do Story principal (mais rápido)
                try:
                    # wdMainTextStory = 1
                    doc_com.StoryRanges(1).Fields.Update()
                except Exception:
                    # Fallback: atualizar todos os campos de Content
                    try:
                        doc_com.Content.Fields.Update()
                    except Exception:
                        pass

                # Salva alterações no arquivo
                doc_com.Save()
            finally:
                try:
                    # Fecha apenas o documento, mantendo o Word aberto
                    doc_com.Close(SaveChanges=False)
                except Exception:
                    pass
            return True
        except Exception as e:
            logger.error(f"Erro geral no update Win32: {e}")
            return False

    def _add_ementa_with_inline_links_to_paragraph(self, p, texto: str):
        try:
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            def add_text_run(t: str):
                if not t:
                    return
                r = p.add_run(t)
                r.font.name = 'Arial MT'
                r.font.size = Pt(8)
            def add_hyperlink_run(display: str, url: str):
                r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
                hl = OxmlElement('w:hyperlink')
                hl.set(qn('r:id'), r_id)
                run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), 'Hyperlink')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Arial MT')
                rFonts.set(qn('w:hAnsi'), 'Arial MT')
                rFonts.set(qn('w:cs'), 'Arial MT')
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '16')
                rPr.append(rStyle)
                rPr.append(rFonts)
                rPr.append(sz)
                run.append(rPr)
                t = OxmlElement('w:t')
                t.text = display
                run.append(t)
                hl.append(run)
                p._p.append(hl)
            pattern = re.compile(r'<(https?://[^>\s]+)>|(https?://\S+)')
            pos = 0
            texto = texto or ''
            for m in pattern.finditer(texto):
                start, end = m.start(), m.end()
                add_text_run(texto[pos:start])
                url = m.group(1) or m.group(2)
                bracketed = m.group(1) is not None
                if bracketed:
                    add_text_run('<')
                add_hyperlink_run(url, url)
                if bracketed:
                    add_text_run('>')
                pos = end
            add_text_run(texto[pos:])
            return p
        except Exception:
            run = p.add_run(texto or '')
            run.font.name = 'Arial MT'
            run.font.size = Pt(8)
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            return p

    def _format_line(self, paragrafo, linha: str):
        parts = linha.split(':', 1)
        if len(parts) == 2 and parts[0].strip() and parts[1] is not None:
            key, value = parts
           
            run_key = paragrafo.add_run(f"{key.strip()}:")
            run_key.bold = True
            run_key.font.name = 'Arial MT'
            run_key.font.size = Pt(8)

            run_val = paragrafo.add_run(f" {value.strip()}")
            run_val.font.name = 'Arial MT'
            run_val.font.size = Pt(8)
        else:
            run = paragrafo.add_run(linha)
            run.font.name = 'Arial MT'
            run.font.size = Pt(8)
            if linha.strip().lower() == 'ementa:':
                run.bold = True

        pf = paragrafo.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE


# Execução: usar pytest para rodar a classe acima, exemplo:
# pytest -q "jt_juris_teste 1.py" -k test_extrair_jt -s --headed

if __name__ == "__main__":
    # Para executar diretamente, sem pytest:
    # python "jt_juris_teste 1.py"
    #
    # Para ver o navegador, defina a variável de ambiente:
    # $env:JT_HEADED="1"
    #
    # Para ativar o log de debug da gravação no DOCX:
    # $env:JT_DEBUG_LOG="1"

    from seleniumbase import SB
    
    # Checa se o modo headed está ativado por variável de ambiente
    headed_mode = os.environ.get("JT_HEADED", "0") == "1"
    
    # Usa o gerenciador de contexto SB() para inicializar corretamente
    with SB(headed=headed_mode, test=True) as sb:
        # Cria instância da classe de teste
        test_case = JTJurisTeste(methodName='test_extrair_jt')
        
        # Injeta o driver do SB na instância do teste
        test_case.driver = sb.driver
        test_case.browser = sb.browser
        
        # Chama setUp para inicializar seletores e outras variáveis
        test_case.selectors = test_case._load_selectors()
        test_case._turma_bookmarks = {}
        test_case._bookmark_id_counter = 1
        test_case.skip_sumario = os.environ.get("JT_SKIP_SUMARIO", "0") == "1"
        test_case.disable_clipboard = os.environ.get("JT_DISABLE_CLIPBOARD", "0") == "1"
        
        # Copia métodos úteis do SB para a instância
        test_case.open = sb.open
        test_case.wait_for_element = sb.wait_for_element
        test_case.find_elements = sb.find_elements
        test_case.find_element = sb.find_element
        
        try:
            # Executa o método de teste
            logger.info("Iniciando execução do teste...")
            test_case.test_extrair_jt()
            logger.info("Teste concluído com sucesso!")
        except Exception as e:
            logger.error(f"Ocorreu um erro durante a execução do teste: {e}")
            logger.error(traceback.format_exc())